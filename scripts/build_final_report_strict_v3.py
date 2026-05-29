# -*- coding: utf-8 -*-
"""城垣杯竞赛成果研究报告构建与格式化 - 严格保留模板格式版 V3"""
import os
import sys
from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
sys.stdout.reconfigure(encoding='utf-8')
TEMPLATE_PATH = Path(r"C:\Users\23902\Desktop\城环杯\附件3 成果研究报告（模板）.docx")
OUTPUT_PATH = Path(r"C:\Users\23902\Desktop\城环杯\附件3 成果研究报告.docx")
TEMP_IMG_DIR = Path(r"C:\Users\23902\Desktop\城环杯\temp_images")
def set_font_run_only(run, font_name="Times New Roman", east_asia_font="宋体"):
    """仅替换字体名称，避免重复创建 w:rFonts 节点以保持 XML 结构严谨"""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), east_asia_font)
def insert_formatted_p(anchor_p, parts, style_name="Normal"):
    # 使用 Normal 样式作为正文段落样式，继承模板中 Normal 的行距、缩进等格式
    p = anchor_p.insert_paragraph_before(style=style_name)
    for text, bold, italic, super_script in parts:
        run = p.add_run(text)
        set_font_run_only(run)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if super_script:
            run.font.superscript = True
    return p
# ── XML 边框与着色辅助函数 ──
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        tag = 'w:{}'.format(edge)
        element = tcBorders.find(qn(tag))
        if element is not None:
            tcBorders.remove(element)
        if edge_data:
            element = OxmlElement(tag)
            tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))
def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is not None:
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)
def apply_three_line_table_style(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
        for cell in row.cells:
            # 单元格边距
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side in ('top', 'bottom', 'left', 'right'):
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'), '120' if side in ('top', 'bottom') else '180')
                m.set(qn('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)
            # 设置边框
            borders = {"left": None, "right": None, "insideV": None, "insideH": None}
            if i == 0:  # 第一行 (表头)
                borders["top"] = {"sz": 12, "val": "single", "color": "000000"}
                borders["bottom"] = {"sz": 6, "val": "single", "color": "000000"}
                set_cell_shading(cell, "F2F2F2")
            elif i == len(table.rows) - 1:  # 最后一行 (底线)
                borders["top"] = None
                borders["bottom"] = {"sz": 12, "val": "single", "color": "000000"}
            else:  # 中间数据行
                borders["top"] = None
                borders["bottom"] = None
            set_cell_border(cell, **borders)
def format_placeholder_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            set_cell_shading(cell, "FAFAFA")
            set_cell_border(
                cell,
                top={"sz": 4, "val": "dashed", "color": "CCCCCC"},
                bottom={"sz": 4, "val": "dashed", "color": "CCCCCC"},
                left={"sz": 4, "val": "dashed", "color": "CCCCCC"},
                right={"sz": 4, "val": "dashed", "color": "CCCCCC"}
            )
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                for run in p.runs:
                    set_font_run_only(run, font_name="Times New Roman", east_asia_font="宋体")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
def format_equation_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=None, bottom=None, left=None, right=None)
            for idx, c in enumerate(row.cells):
                for p in c.paragraphs:
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    if idx == 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in p.runs:
                        set_font_run_only(run, font_name="Times New Roman", east_asia_font="宋体")
def insert_table_before(anchor_p, rows, cols, style_name="Table Grid"):
    """在 anchor_p 之前插入表格，确保 XML 顺序完全正确"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = style_name
    # 插入到 anchor_p 的前面
    p_parent = anchor_p._element.getparent()
    p_parent.insert(p_parent.index(anchor_p._element), table._element)
    return table
def insert_img_before(anchor_p, img_path, caption_text, width_inches=4.8):
    """在 anchor_p 之前依次插入图片段落和居中的图题段落"""
    p_img = anchor_p.insert_paragraph_before()
    p_img.style = "Normal"
    p_img.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.line_spacing = 1.0
    p_img.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_img.paragraph_format.space_before = Pt(12)
    p_img.paragraph_format.space_after = Pt(6)
    p_img.paragraph_format.first_line_indent = Pt(0)
    run_img = p_img.add_run()
    run_img.add_picture(str(img_path), width=Inches(width_inches))
    p_cap = anchor_p.insert_paragraph_before()
    p_cap.style = "Normal"
    p_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.line_spacing = Pt(20)
    p_cap.paragraph_format.space_before = Pt(0)
    p_cap.paragraph_format.space_after = Pt(12)
    p_cap.paragraph_format.first_line_indent = Pt(0)
    run_cap = p_cap.add_run(caption_text)
    set_font_run_only(run_cap, font_name="Times New Roman", east_asia_font="黑体")
    run_cap.bold = True
    run_cap.font.size = Pt(10.5)
    return p_cap
def insert_equation_before(anchor_p, eq_text, eq_num_text):
    """在 anchor_p 之前插入公式表"""
    table = insert_table_before(anchor_p, rows=1, cols=2)
    # 设为无边框
    borders_xml = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="none"/><w:left w:val="none"/>'
        '<w:bottom w:val="none"/><w:right w:val="none"/>'
        '<w:insideH w:val="none"/><w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    table._element.tblPr.append(borders_xml)
    row = table.rows[0]
    cell_eq = row.cells[0]
    cell_num = row.cells[1]
    cell_eq.width = Inches(5.5)
    cell_num.width = Inches(1.0)
    p_eq = cell_eq.paragraphs[0]
    p_eq.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eq.paragraph_format.line_spacing = Pt(20)
    p_eq.paragraph_format.space_before = Pt(6)
    p_eq.paragraph_format.space_after = Pt(6)
    p_eq.paragraph_format.first_line_indent = Pt(0)
    run_eq = p_eq.add_run(eq_text)
    set_font_run_only(run_eq)
    run_eq.italic = True
    p_num = cell_num.paragraphs[0]
    p_num.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_num.paragraph_format.line_spacing = Pt(20)
    p_num.paragraph_format.space_before = Pt(6)
    p_num.paragraph_format.space_after = Pt(6)
    p_num.paragraph_format.first_line_indent = Pt(0)
    run_num = p_num.add_run(eq_num_text)
    set_font_run_only(run_num)
    return table
def find_paragraph_by_text(doc, text_fragment):
    for idx, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return idx
    return -1
def replace_section_content_before(doc, subheading_text, next_element_text, populate_func):
    idx_sub = find_paragraph_by_text(doc, subheading_text)
    if idx_sub == -1:
        print(f"未找到小标题：{subheading_text}")
        return
    idx_next = find_paragraph_by_text(doc, next_element_text)
    if idx_next == -1:
        print(f"未找到下一段：{next_element_text}")
        return
    # 仅删除该区域内的正文段落，如遇分节符则仅清除文字以保留格式
    p_to_del = [doc.paragraphs[i] for i in range(idx_sub + 1, idx_next)]
    for p in p_to_del:
        if len(p._element.xpath('.//w:sectPr')) > 0:
            p.text = "" # 保留分节符
        else:
            p._element.getparent().remove(p._element)
    idx_next = find_paragraph_by_text(doc, next_element_text)
    p_next = doc.paragraphs[idx_next]
    populate_func(p_next)
# ── 加载文档 ──
print(f"加载模板文档: {TEMPLATE_PATH}")
doc = Document(str(TEMPLATE_PATH))
# ── 1. 填入封面参赛信息（修改文字，严格保留原有字体与字号样式） ──
print("填入封面参赛信息（仅修改文本，保持原有字体与字号样式不变）...")
def update_cover_paragraph(p, replacement_text, clear_others=True):
    if len(p.runs) > 0:
        p.runs[0].text = replacement_text
        if clear_others:
            for r in p.runs[1:]:
                r.text = ""
    else:
        p.text = replacement_text
    # 强制将行距设为多倍行距（1.25倍），防止长文本折行时发生文字重叠
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
for idx, p in enumerate(doc.paragraphs[:30]):
    text = p.text.strip()
    if "题目（35个汉字以内" in text:
        update_cover_paragraph(p, "基于大模型与多模态AI的城市更新\n空间设计智能推演系统")
    elif "参赛编号：" in text:
        update_cover_paragraph(p, "参赛编号：F221")
    elif "参 赛 人：" in text:
        update_cover_paragraph(p, "参 赛 人：陈礼冲、刘旭东")
    elif "指 导 人：" in text:
        update_cover_paragraph(p, "指 导 人：李冰心、崔诚慧")
    elif "工作单位：" in text:
        update_cover_paragraph(p, "工作单位：吉林建筑大学")
    elif "报名主题：" in text:
        update_cover_paragraph(p, "报名主题：主题二：面向高质量发展的城市治理")
    elif "研究议题：" in text or "研究课题：" in text:
        label = "研究课题：" if "研究课题：" in text else "研究议题："
        val = "课题6：城市体检与城市更新" if "研究课题" in text else "议题6：城市体检与城市更新"
        update_cover_paragraph(p, f"{label}{val}")
    elif "技术关键词：" in text or "技术路径：" in text or "关键技术：" in text:
        if "技术路径：" in text:
            update_cover_paragraph(p, "技术路径：技术路径一：AI建模、检测、生物相似性/人机融合建模与分析（请选一）")
        elif "关键技术：" in text:
            update_cover_paragraph(p, "关键技术（三选一）：大语言模型（必填）、多主体博弈决策（选填）、城市体检诊断（选填）")
        else:
            update_cover_paragraph(p, "技术关键词：大语言模型（必填）、多主体博弈决策（选填）、城市体检诊断（选填）")
    elif "介绍参赛团队的研究背景" in text:
        update_cover_paragraph(p, "本参赛团队成员均来自吉林建筑大学城乡规划专业，长期致力于数字城乡规划与智能空间决策支持系统研究。团队在多源空间大数据分析、遥感与计算机视觉处理、大语言模型与多智能体博弈系统在规划中的应用等方面具备深厚积淀。成员曾深度参与多项吉林省及长春市历史文化保护街区微更新实证项目，拥有丰富的数字平台开发与诊断分析经验。")
    elif "字数要求：100-200字" in text:
        update_cover_paragraph(p, "")
# ── 2. 清除模板格式指南 〇 （保留任何含分节符的段落） ──
print("清除模板格式指南 〇...")
body = doc._body._body
start_idx = -1
end_idx = -1
for idx, child in enumerate(body):
    text = ''
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        text = child.text or ''
    elif tag == 'tbl':
        text = 'TABLE: ' + ''.join([t.text for t in child.xpath('.//w:t')])
    if '〇、正文格式模板' in text:
        start_idx = idx
    if '一、研究问题' in text:
        end_idx = idx
        break
if start_idx != -1 and end_idx != -1:
    # 倒序删除以避免索引错位
    for idx in range(end_idx - 1, start_idx - 1, -1):
        child = body[idx]
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            p = docx.text.paragraph.Paragraph(child, doc)
            if len(p._element.xpath('.//w:sectPr')) > 0:
                p.text = "" # 含有分节符，保留节点以维持页面和分节属性，仅清除文字
            else:
                body.remove(child) # 无分节符，直接安全删除
        else:
            body.remove(child) # 示例表格等直接删除
# ── 3. 逐步注入正文（继承模板格式） ──
# 第一章第一节
def populate_c1_s1(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("我国城市化发展已由“大拆大建”的增量扩张阶段全面转向“存量提质”的高质量发展阶段。以老旧小区改造、历史风貌协调街区整治和低效用地腾退为核心的城市微更新，已成为转变城市发展方式、提升人居环境品质的重要抓手。然而，在城市微更新的学术研究与具体规划实践中，仍面临三大长期瓶颈：第一，在现状诊断层面，传统的城市体检高度依赖规划专家的主观定性判断与粗粒度普查，缺乏可量化、地块级、多源数据融合的更新潜力诊断模型；第二，在更新决策层面，微更新涉及政府、开发商、居民等多方利益主体，各方诉求（历史保护、经济回报、生活便利）冲突激烈，缺乏科学的动态博弈与共识达成平台；第三，在规划表达与方案深化层面，传统计算机辅助设计与地理信息系统（Computer Aided Design / Geographic Information System，CAD/GIS）手工制图周期长、效率低，且直接应用生成式人工智能（Artificial Intelligence Generated Content，AIGC）进行规划绘图时存在严重的空间位置幻觉与地理控制偏差。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("针对上述痛点，本研究依托长春伪满皇宫周边170.4公顷的典型历史文化与工业协调更新片区，基于大语言模型（Large Language Model，LLM）与多模态人工智能技术，自主研发了“城市更新空间设计智能推演系统（ultimateDESIGN）”。该系统深度融合了四大AI核心技术：一是基于SegFormer语义分割神经网络的多源空间感知与自动化城市体检，将街景图片全自动转化为绿视率等定量指标；二是LLM驱动的多主体利益博弈闭环协商机制，将文本推理与空间参数效用函数深度绑定；三是基于检索增强生成（Retrieval-Augmented Generation，RAG）的控规合规性智能审计，实现248条法规的秒级自动比对；四是首创「矢量-光栅-ControlNet」刚性对齐管线，从根本上消除AIGC制图的地理空间幻觉。上述四大引擎构建了“诊断-博弈-生成-校验”的闭环循证工作流，为探索存量规划的民主决策与高精度制图设计提供了全新的方法论支撑。", False, False, False)
    ])
replace_section_content_before(doc, "1．研究背景及目的意义", "2．研究目标及拟解决的问题", populate_c1_s1)
# 第一章第二节
def populate_c1_s2(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本项目的总体目标是针对城市微更新中“诊断粗糙”、“博弈断裂”和“制图幻觉”三大痛点，构建一套基于多模态大模型的城市微更新全生命周期数字化辅助系统。具体目标包括：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）构建基于层次分析法（Analytic Hierarchy Process，AHP）与微更新潜力指数（Micro-Planning Index，", False, False, False),
        ("MPI", False, True, False),
        ("）的地块级多维定量评估模型。依托SegFormer语义分割神经网络批量解析1,788张街景图片，将绿视率、天空开阔度等视觉风貌特征全自动转化为定量指标，输入AHP-MPI模型，对研究范围内719栋现状建筑及相应地块的更新紧迫性与潜力进行科学排序；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）建立LLM驱动的“居民-开发商-规划师”三主体利益博弈协商沙盘。将大语言模型的文本推理能力与空间参数效用函数（容积率、绿地率、限高等）深度绑定，基于关键词命中退避机制动态评估多方满意度，当检测到满意度低于共识线时触发冲突预警并由AI规划师智能体根据RAG检索到的政策进行妥协式方案修正，直至三方满意度收敛；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）首创「矢量-光栅-ControlNet」空间约束AI制图管线。将GIS矢量规划红线与道路网转为光栅掩膜，作为ControlNet的条件约束输入Stable Diffusion模型，强制AI生成的意向效果图严格贴合规划边界，从根本上消除AIGC制图的地理空间幻觉，实现26张A3标准规划图纸的全自动多进程并行编译与专业排版；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（4）以长春市伪满皇宫博物院周边170.4公顷街区为实证应用，验证系统在实际存量城市微更新项目中的可用性与可迁移性。", False, False, False)
    ])
replace_section_content_before(doc, "2．研究目标及拟解决的问题", "二、研究方法", populate_c1_s2)
# 第二章第一节
def populate_c2_s1(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本系统依托计算机视觉、多智能体协作、空间拓扑计算以及大语言模型，建立了一套人机协同的城市微更新分阶段决策支持体系。核心方法与理论依据如下：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）空间网络拓扑与步行可达性计算：参考丁梦月的街道步行空间分析方法", False, False, False),
        ("[1]", False, False, True),
        ("，系统依托OSMnx与NetworkX库对研究范围内74段核心路网进行拓扑建模与测算，得出全局整合度（", False, False, False),
        ("Integration", False, True, False),
        ("）与穿行度（", False, False, False),
        ("Choice", False, True, False),
        ("），定量刻画道路网络在微更新中的空间织补与慢行可步行网络连通需求；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）深度学习街景测度与风貌基因图谱提取：参考尧馨雅的街道风貌基因识别方法", False, False, False),
        ("[2]", False, False, True),
        ("，采用SegFormer（基于Transformer架构的语义分割模型）深度学习网络，对伪满皇宫周边447个采样点共1,788张实景影像进行像素级语义分割，自动测度街景绿视率（Green View Index，", False, False, False),
        ("GVI", False, True, False),
        (" = 8.7%）与天空开阔度（Sky View Factor，", False, False, False),
        ("SVF", False, True, False),
        ("），以量化环境现状品质；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）历史肌理延续与AIGC空间生成约束：参考赵卉的历史肌理数字化城市设计方法", False, False, False),
        ("[3]", False, False, True),
        ("，引入「矢量-光栅-ControlNet」刚性约束AI制图管线，将规划红线与用地矢量转换为带国标色值的光栅图输入Stable Diffusion模型，消除AIGC生成式设计意向图的结构漂移与形态幻觉；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（4）多主体博弈协同与社会利益调停：参考卢文正的社会空间更新理论", False, False, False),
        ("[4]", False, False, True),
        ("，系统设立“居民-开发商-政府”三方智能体，模拟在利益协商中的角色对立，利用大语言模型（deepseek-v4-flash）进行多轮博弈协商，并通过文本语义命中机制评估满意度走势；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（5）时空数据库设计与多源数据汇聚：参考张峰的空间信息资源规划模型", False, False, False),
        ("[5]", False, False, True),
        ("，设计了统一的时空数据库与Stage Data Bus总线，实现GIS矢量、时空轨迹、文本等多源异构数据的全周期互锁与状态监测；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（6）数字孪生城市韧性与防灾分析：参考张国政的数字孪生韧性提升路径", False, False, False),
        ("[6]", False, False, True),
        ("，系统基于Pydeck三维数字底盘，实现伪满皇宫周边719栋建筑现状白模与防灾漫溢模拟，开展多维度空间韧性诊断；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（7）旧工业区改造与步行活力测度：参考梁汉雄的可步行性要素评价方法", False, False, False),
        ("[7]", False, False, True),
        ("，建立基于核密度估计（KDE）的POI服务网格，量化居民日常出行便捷度与旧厂区更新微循环的慢行品质潜力；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（8）旧城有机更新机制与合规审计规范：参考方可的北京旧城居住区有机更新途径", False, False, False),
        ("[8]", False, False, True),
        ("，构建RAG检索增强生成的法规比对库，进行法定退界、建筑高度等容积率控规指标的动态校验，越界自动判定违规红牌。", False, False, False)
    ])
    img_path = TEMP_IMG_DIR / "urban_rural_planning_mindmap.png"
    if img_path.exists():
        insert_img_before(p, img_path, "图2-1  城乡规划设计规范与多模态AI智能体知识体系图", width_inches=5.0)
replace_section_content_before(doc, "1．研究方法及理论依据", "2．技术路线及关键技术", populate_c2_s1)
# 第二章第二节
def populate_c2_s2(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本系统技术路线分为数据汇聚层、计算引擎层、推演策略层、以及成果表达层四个部分：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("数据层汇总了包括建筑基底在内的6类GeoJSON矢量数据和15份包含街景、POI在内的CSV感知数据，统一纠正为本地投影；计算层基于AHP和空间句法，对719栋建筑地块进行定量化多维问题诊断；策略层依托大模型（deepseek-v4-flash与deepseek-v4-pro），进行三主体博弈协商、Zoning Compliance RAG合规性实时审查、规划说明文本动态编译，并由Stage Transition Agent和Copilot sidebar提供跳转导航；表达层利用Matplotlib与Pillow组合构建了“三层排版流水线”，自动拼装包含比例尺、图例、图签以及大模型动态设计说明的标准A3规划大图，并支持多进程并行批量图纸编译。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("系统的四大关键技术包括：一是「矢量-光栅-ControlNet」空间约束AI制图管线，将用地、道路等矢量红线转换为带国标色值的光栅图作为ControlNet of Stable Diffusion的输入，彻底消除了生成式AI在规划底图上的空间位置漂移与形态幻觉；二是Data-to-Text大模型动态图例说明合成技术，在A3底板封装阶段，大模型（deepseek-v4-flash）读取GIS实测指标，自动生成设计说明文字，确保出图指标与真实GIS物理数据严格绑定；三是多进程并行处理技术，充分调动多核CPU进行26张图纸的并行渲染，极大缩短图册编译时间；四是Zoning RAG 控规合规自动比对审计，通过大语言模型检索本地规范文本，在Streamlit前端对超出红线的违规指标高亮违规警告。", False, False, False)
    ])
    img_arch = TEMP_IMG_DIR / "system_architecture_mindmap.png"
    if img_arch.exists():
        insert_img_before(p, img_arch, "图2-2  ultimateDESIGN 决策支持系统四层架构设计图", width_inches=5.0)
    img_flow = TEMP_IMG_DIR / "workflow_flowchart.png"
    if img_flow.exists():
        insert_img_before(p, img_flow, "图2-3  ultimateDESIGN 17阶段全生命周期设计推演工作流图", width_inches=5.0)
replace_section_content_before(doc, "2．技术路线及关键技术", "三、数据说明", populate_c2_s2)
# 第三章第一节 (数据清单表格 3-1)
def populate_c3_s1(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("项目实证研究建立在多源、高精度空间数据库之上。所涉及的主要数据资产包括以下8大类：", False, False, False)
    ])
    # 表名在表格上方
    p_cap = p.insert_paragraph_before()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.line_spacing = Pt(20)
    p_cap.paragraph_format.space_before = Pt(12)
    p_cap.paragraph_format.space_after = Pt(6)
    p_cap.paragraph_format.first_line_indent = Pt(0)
    run_cap = p_cap.add_run("系统多源数据资产清单  表3-1")
    set_font_run_only(run_cap, font_name="Times New Roman", east_asia_font="黑体")
    run_cap.bold = True
    run_cap.font.size = Pt(10.5)
    table = insert_table_before(p, rows=9, cols=5)
    hdr_cells = table.rows[0].cells
    headers = ['序号', '数据名称', '数据格式', '数据规模', '在模型设计中的作用与来源']
    for idx, val in enumerate(headers):
        p_cell = hdr_cells[idx].paragraphs[0]
        p_cell.text = ""
        run = p_cell.add_run(val)
        set_font_run_only(run, font_name="Times New Roman", east_asia_font="黑体")
        run.bold = True
    data_rows = [
        ("1", "研究范围红线", "GeoJSON", "170.4 公顷", "确立规划空间边界约束，自绘"),
        ("2", "现状建筑基底", "GeoJSON", "110,289 栋", "提取层高、建筑面积以计算容积率及建筑密度，OSM"),
        ("3", "道路网络", "GeoJSON", "1,062 段核心路段", "用于空间句法拓扑可达性及网络穿行度计算，OSM"),
        ("4", "现状用地分类", "GeoJSON", "1,026 宗地块", "核查现状与控规用地占比，计算现状绿地率，自绘"),
        ("5", "兴趣点（Point of Interest，POI）", "CSV", "411 条", "计算社会设施配套需求与服务核密度，百度API"),
        ("6", "街景影像", "JPG", "1,788 张", "计算各街区环境绿视率(GVI)及天空开阔度，百度街景"),
        ("7", "舆情文本", "CSV", "207 条", "情感分析模型获取公众对历史街区更新的社会诉求，新浪微博"),
        ("8", "政策保护规章", "PDF/JSON", "248 个向量分块", "输入RAG知识库，用于导则生成与合规审查，自然资源部/住建部")
    ]
    for i, row in enumerate(data_rows):
        row_cells = table.rows[i+1].cells
        for j, val in enumerate(row):
            p_cell = row_cells[j].paragraphs[0]
            p_cell.text = ""
            run = p_cell.add_run(val)
            set_font_run_only(run, font_name="Times New Roman", east_asia_font="宋体")
    apply_three_line_table_style(table)
replace_section_content_before(doc, "1．数据内容及类型", "2．数据预处理技术与成果", populate_c3_s1)
# 第三章第二节
def populate_c3_s2(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("为了保证各项指标测算及AI空间绘图定位的物理精度，我们执行了严格的数据预处理流程：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）本地投影纠偏（EPSG:32652）：多源获取的空间数据往往存在WGS-84、GCJ-02、BD-09等多种空间坐标系偏差。本系统首先开发了坐标系批量转换工具（`geo_transform.py`），将所有坐标统一转化为地理坐标系，并选择以长春本地的高斯克吕格投影（EPSG:32652，北京54/3度分带第42带，中央经线126E，或UTM Zone 52N）进行高精度平面投影。这有效避免了Web墨卡托（EPSG:3857）由于投影拉伸在中高纬度地区带来的面积计算偏差（高纬度形变率高达93%），将研究范围精确锚定在170.4公顷。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）多源属性空间挂接：利用空间拓扑Join技术，将411条POI配套兴趣点、1,788张街景图片以及房屋年代、价格等非空间CSV属性，按物理空间质心精确挂接到相应的地块及719栋建筑矢量基底中，构建六维度的空间上下文属性库。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）街景绿视率(GVI)计算机视觉测度：针对下载的1,788张百度实地四方向全景街景，使用训练好的SegFormer", False, False, False),
        ("[2]", False, False, True),
        ("（ADE20K）图像语义分割神经网络，提取街景中的绿化植物植被像素比率，输出447个物理样点的全域绿视率平均值（", False, False, False),
        ("GVI", False, True, False),
        (" = 8.7%），识别出历史街区内部绿化品质严重不足的物理事实。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（4）规章切块与向量嵌入：提取长春市及伪满皇宫保护管理规划的7份核心PDF文件，使用递归字符切片法将其切分为248个包含上下文重叠的语义块，采用BGE（BAAI/bge-large-zh-v1.5）中文大模型提取高维语义嵌入向量，导入到系统的轻量级本地向量数据库中，构建RAG知识库。", False, False, False)
    ])
    img_pipe = TEMP_IMG_DIR / "data_pipeline_mindmap.png"
    if img_pipe.exists():
        insert_img_before(p, img_pipe, "图3-1  多源空间数据处理与特征提取数据管线图", width_inches=5.0)
replace_section_content_before(doc, "2．数据预处理技术与成果", "四、模型算法", populate_c3_s2)
# 第四章第一节 (包含公式4-1, 4-2, 4-3)
def populate_c4_s1(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本系统在数据底座之上，设计并实现了三套核心规划分析与智能推演算法：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）AHP-", False, False, False),
        ("MPI", False, True, False),
        (" 空间微更新潜力指数计算模型：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("微更新潜力测度是确定老旧街区体检“病灶地块”的核心。地块的更新潜力指数由空间潜力（", False, False, False),
        ("S", False, True, False),
        ("）、社会配套需求（", False, False, False),
        ("D", False, True, False),
        ("）以及环境现状品质（", False, False, False),
        ("E", False, True, False),
        ("）三维度综合加权得出，数学公式如下：", False, False, False)
    ])
    insert_equation_before(p, "MPI_i = (0.4 * S_i + 0.3 * D_i + 0.3 * (1.0 - E_i)) * 100", "(4-1)")
    insert_formatted_p(p, [
        ("其中，", False, False, False),
        ("MPI_i", False, True, False),
        (" 表示第 ", False, False, False),
        ("i", False, True, False),
        (" 栋现状建筑的微更新潜力指数；", False, False, False),
        ("S_i", False, True, False),
        (" 表示其空间潜力（通过地块基底面积、建筑层数倒数归一化得到）；", False, False, False),
        ("D_i", False, True, False),
        (" 表示其设施服务配套的社会需求（基于411个POI按150米搜寻半径做核密度估算得到）；", False, False, False),
        ("E_i", False, True, False),
        (" 表示采样样点的街景绿视率", False, False, False),
        ("GVI", False, True, False),
        ("均值（由SegFormer语义分割得出）。其数学内涵在于：利用 (1.0 - ", False, False, False),
        ("E_i", False, True, False),
        (") 表征当前绿化品质越差，其微更新的迫切性及环境改善需求度越高。三维度权重采用层次分析法（AHP）判定矩阵运算划分（空间 ", False, False, False),
        ("S", False, True, False),
        (": 0.4，社会 ", False, False, False),
        ("D", False, True, False),
        (": 0.3，环境 ", False, False, False),
        ("E", False, True, False),
        (": 0.3，通过一致性比例 ", False, False, False),
        ("CR", False, True, False),
        (" < 0.1 检验），保证了空间品质精准诊断的数学严谨性。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）多主体协同协商的 LLM 满意度效用评估模型：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("为了协调老旧社区微更新中多方利益冲突，系统设计了多智能体协商会话沙盘。预置了居民代表、开发商以及规划师（政府）三大角色，并利用大语言模型对多轮发言进行语义理解与效用自动评分。大语言模型输出的满意度分值由其内置的评分提示词及JSON返回矩阵确定，为了确保评估的鲁棒性，系统设计了确定性的关键词命中退避机制：", False, False, False)
    ])
    insert_equation_before(p, "S_role = min(100, 50 + 7 * sum( [1 for word in K_role if word in DialogueText] ))", "(4-2)")
    insert_formatted_p(p, [
        ("其中，", False, False, False),
        ("S_role", False, True, False),
        (" 表示角色满意度；", False, False, False),
        ("K_role", False, True, False),
        (" 表示各方的核心关键词诉求集：居民代表 ", False, False, False),
        ("K_res", False, True, False),
        (" = [“绿”, “公园”, “配套”, “社区”, “医院”, “菜市”, “养老”, “口袋”]；开发商 ", False, False, False),
        ("K_dev", False, True, False),
        (" = [“容积率”, “收益”, “文旅”, “商业”, “民宿”, “运营”, “产业”, “投资”, “回报”]；规划师 ", False, False, False),
        ("K_gov", False, True, False),
        (" = [“历史保护”, “紫线”, “限高”, “合规”, “风貌”, “条例”, “保护区”]；", False, False, False),
        ("DialogueText", False, True, False),
        (" 表示博弈协商过程中的发言对话文本。三方初始满意度设为50分，每发言命中一个核心诉求词，对应满意度得分累加7分，封顶为100分。系统判定整体规划共识达成的条件为三方满意度底线均大于等于60分（即 min(", False, False, False),
        ("S_res", False, True, False),
        (", ", False, False, False),
        ("S_dev", False, True, False),
        (", ", False, False, False),
        ("S_gov", False, True, False),
        (") >= 60），当判定未通过时触发黄牌冲突警告并促使智能体重新调整妥协方案。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）Zoning Compliance RAG 与 GIS 控规实时合规校验算法：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("系统通过时空几何叠加计算，实时校验地块容积率（", False, False, False),
        ("FAR", False, True, False),
        ("）与建筑密度（", False, False, False),
        ("Density", False, True, False),
        ("），公式如下：", False, False, False)
    ])
    insert_equation_before(p, "FAR = sum( Floor_i * Area_i ) / Area_land,  Density = sum( Area_footprint ) / Area_land", "(4-3)")
    insert_formatted_p(p, [
        ("其中，", False, False, False),
        ("FAR", False, True, False),
        (" 表示容积率；", False, False, False),
        ("Density", False, True, False),
        (" 表示建筑密度；", False, False, False),
        ("Floor_i", False, True, False),
        (" 和 ", False, False, False),
        ("Area_i", False, True, False),
        (" 分别表示第 ", False, False, False),
        ("i", False, True, False),
        (" 栋现状建筑的层数和基底面积；", False, False, False),
        ("Area_land", False, True, False),
        (" 表示地块红线面积；", False, False, False),
        ("Area_footprint", False, True, False),
        (" 表示现状建筑基底总面积。除了经典的数值比对，系统在后端将开发强度方案输入RAG本地法规向量库进行余弦相似度匹配，检索最相关的3条法规条文（匹配阈值 >= 0.65），交由大模型（deepseek-v4-flash）进行文字层面的越界红线审计，输出结构化的合规判定 JSON 数组，当检测到强红线违规时自动高亮触发红牌告警。", False, False, False)
    ])
    img_neg = TEMP_IMG_DIR / "agent_negotiation_flowchart.png"
    if img_neg.exists():
        insert_img_before(p, img_neg, "图4-1  三方主体智能协同决策博弈协商图", width_inches=5.0)
    img_param = TEMP_IMG_DIR / "technology_parameters_knowledge_graph.png"
    if img_param.exists():
        insert_img_before(p, img_param, "图4-2  规划管控系统参数与多模态知识图谱映射图", width_inches=5.0)
replace_section_content_before(doc, "1．模型算法流程及相关数学公式", "2．模型算法相关支撑技术", populate_c4_s1)
# 第四章第二节
def populate_c4_s2(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本系统完全基于纯 Python 开发，开发与部署栈包括以下支撑技术：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）基础语言与平台：采用 Python 3.12 核心，前端交互与可视化采用 Streamlit 1.55 全栈框架；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）GIS 空间几何库：核心基于 GeoPandas 1.0、Shapely 2.0 以及 PyProj 3.6 进行投影和缓冲区拓扑运算；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）网络分析库：基于 NetworkX 3.2 和 OSMnx 1.9 对74段核心规划道路网的全局整合度进行连通性拓扑图论分析；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（4）可视化组件：引入 Plotly 5.24 进行 3D WebGL 交互式建筑高度基底渲染及博弈雷达图绘制；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（5）大模型及向量搜索：后台调用统一的 DeepSeek-V4 (Pro) API 大语言模型，利用 NumPy 和 Scikit-learn 对248个BGE语义向量进行基于余弦相似度的余弦距离检索。图纸编译调用 Python Multiprocessing 实现多进程并行，大幅优化系统出图效率。", False, False, False)
    ])
    img_rag = TEMP_IMG_DIR / "rag_compliance_flowchart.png"
    if img_rag.exists():
        insert_img_before(p, img_rag, "图4-3  Zoning RAG 法规合规审查流线图", width_inches=5.0)
    img_sd = TEMP_IMG_DIR / "sd_controlnet_flowchart.png"
    if img_sd.exists():
        insert_img_before(p, img_sd, "图4-4  矢量-光栅-ControlNet 空间刚性对齐制图管线图", width_inches=5.0)
replace_section_content_before(doc, "2．模型算法相关支撑技术", "五、实践案例", populate_c4_s2)
# 第五章第一节 (合规对照表 5-1)
def populate_c5_s1(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本实证案例运行于“长春伪满皇宫周边街区微更新支持平台”系统。研究以长春市宽城区伪满皇宫博物院周边170.4公顷街区为对象，范围内包含现状建筑719栋，多为低层和多层历史风貌建筑。整个平台应用过程包含以下三个关键步骤，便于开展全流程功能监测与截图记录：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）第一步：数据载入与多维体检诊断（对应平台“数据准备”与“问题诊断”页面）。", False, True, False),
        (" 决策者首先在数据准备页面上传GeoJSON边界及建筑矢量底图，系统实时验证数据完备性。随后进入问题诊断页面，系统调用AHP-MPI算法对这170.4公顷范围内的719栋建筑进行实测指标提取，并与法定控规上限进行校对，生成城市体检合规诊断对照表如下：", False, False, False)
    ])
    # 表名在表格上方
    p_cap = p.insert_paragraph_before()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.line_spacing = Pt(20)
    p_cap.paragraph_format.space_before = Pt(12)
    p_cap.paragraph_format.space_after = Pt(6)
    p_cap.paragraph_format.first_line_indent = Pt(0)
    run_cap = p_cap.add_run("长春伪满皇宫周边历史街区城市体检合规诊断对照表  表5-1")
    set_font_run_only(run_cap, font_name="Times New Roman", east_asia_font="黑体")
    run_cap.bold = True
    run_cap.font.size = Pt(10.5)
    table_comp = insert_table_before(p, rows=5, cols=4)
    hdr = table_comp.rows[0].cells
    headers = ['体检指标', '实测数值', '法定控规标准', '合规诊断判定']
    for idx, val in enumerate(headers):
        p_cell = hdr[idx].paragraphs[0]
        p_cell.text = ""
        run = p_cell.add_run(val)
        set_font_run_only(run, font_name="Times New Roman", east_asia_font="黑体")
        run.bold = True
    comp_data = [
        ("容积率 (FAR)", "1.13", "≤ 1.40", "✅ 合规达标"),
        ("建筑密度", "30.0%", "≤ 35.0%", "✅ 合规达标"),
        ("绿地率 (GAR)", "2.9%", "≥ 25.0%", "❌ 严重违规 (偏低)"),
        ("最高建筑高度", "59.5 m", "≤ 18.0 m", "⚠️ 局部超高 (站前区溢出)")
    ]
    for i, row in enumerate(comp_data):
        r_cells = table_comp.rows[i+1].cells
        for j, val in enumerate(row):
            p_cell = r_cells[j].paragraphs[0]
            p_cell.text = ""
            run = p_cell.add_run(val)
            set_font_run_only(run, font_name="Times New Roman", east_asia_font="宋体")
    apply_three_line_table_style(table_comp)
    insert_formatted_p(p, [
        ("由体检诊断结果（表5-1）可知，虽然该区域的容积率（1.13）和建筑密度（30.0%）控制较好，但绿地率仅为2.9%，远低于25%的法定标准，这暴露出绿化空间严重不足的短板。同时，外围个别高层建筑高度（最大达59.5m）存在高度超标，破坏了历史文化街区的整体天际线视廊。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）第二步：多主体博弈决策协商（对应平台“设计策略”博弈页面）。", False, True, False),
        (" 针对体检中暴露的高度和绿地问题，系统在设计策略页面模拟了三方利益博弈协商。开发商为了高投资回报，强烈要求将限高放宽至36m，在对话流中多次提及“投资回报、文旅产业”，满意度冲至80分，但规划局坚守“历史保护区、紫线控制、合规条例”底线，满意度为65分，居民代表则强烈反对高层压迫，要求增加口袋公园和生活配套，满意度得分仅为45分。由于 min(", False, False, False),
        ("S_res", False, True, False),
        (", ", False, False, False),
        ("S_dev", False, True, False),
        (", ", False, False, False),
        ("S_gov", False, True, False),
        (") = 45 < 60，触发利益冲突警报。在此警告下，系统引导开发商退让，调整方案将核心区限高控制在18m以内，并在站前区兼容商办，同时配套建设3个口袋公园。经过第二轮博弈，三方满意度得分提升至66分、73分和68分，全部越过60分共识线，系统成功达成更新策略共识。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）第三步：合规自动审计与成果图册编译输出（对应平台“城市设计导则”与“成果表达”页面）。", False, True, False),
        (" 在达成策略共识后，决策者可在城市设计导则页面触发Zoning RAG合规实时校验，生成红牌/黄牌警告面板。最后，在成果表达页面一键触发多进程并行编译引擎，全自动组装比例尺、图例、动态设计说明等标准A3规划大图，完成成果图纸和报告文书的最终输出。", False, False, False)
    ])
replace_section_content_before(doc, "1．模型应用实证及结果解读", "2．模型应用案例可视化表达", populate_c5_s1)
# 第五章第二节
def populate_c5_s2(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("系统的可视化模块为规划方案决策提供了直观、实时的图像表现力：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）平台交互界面与博弈雷达图：在Streamlit页面中，决策者可直观阅读三智能体动态对话流，Plotly雷达图展示三方满意度的实时演进（图5-1），并弹出合规校验红牌告警（图5-2）。", False, False, False)
    ])
    img_radar = TEMP_IMG_DIR / "fig_radar.png"
    if img_radar.exists():
        insert_img_before(p, img_radar, "图5-1  协同博弈阶段多方满意度雷达图及共识判定界面", width_inches=4.8)
    img_comp = TEMP_IMG_DIR / "fig_compliance.png"
    if img_comp.exists():
        insert_img_before(p, img_comp, "图5-2  Streamlit 平台 GIS 控规实时合规性校验告警面板", width_inches=4.8)
    insert_formatted_p(p, [
        ("（2）三维数字底座现状白模与溢出模拟：系统主页集成了基于Pydeck的三维数字twin底盘，直观展现了伪满皇宫周边719栋建筑的风貌层高现状，并具备降雨漫溢防灾的物理推演诊断能力（图5-3）。", False, False, False)
    ])
    img_3d = TEMP_IMG_DIR / "fig_3d.png"
    if img_3d.exists():
        insert_img_before(p, img_3d, "图5-3  三维数字孪生现状风貌与漫溢模拟底座", width_inches=4.8)
    insert_formatted_p(p, [
        ("（3）A3标准成果图册渲染与视觉方案设计：系统全自动编译并导出了高清标准A3图册。图纸右侧图例框下半部分排版了由大模型根据实测指标生成的规划说明（图5-4），并利用“矢量-光栅-ControlNet”制图管线完成绿地景观推演设计（图5-5）。", False, False, False)
    ])
    img_dr004 = TEMP_IMG_DIR / "fig_004.png"
    if img_dr004.exists():
        insert_img_before(p, img_dr004, "图5-4  系统自动编译导出的 A3 成果图纸：DR-004 现状区位图 (自绘图例与动态设计说明)", width_inches=5.2)
    img_land = TEMP_IMG_DIR / "unified_landscape_mindmap.png"
    if img_land.exists():
        insert_img_before(p, img_land, "图5-5  重点更新单元绿地景观系统设计推演图", width_inches=5.2)
replace_section_content_before(doc, "2．模型应用案例可视化表达", "六、研究总结", populate_c5_s2)
# 第六章第一节
def populate_c6_s1(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本系统设计的特点和创新性突出表现在以下四个维度：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）四大AI引擎深度耦合的闭环决策流：首次将基于SegFormer神经网络的“AHP-MPI体检诊断模型”、LLM驱动的“多主体博弈决策协商模型”、基于BGE向量检索的“RAG政策合规校验模型”与基于ControlNet的“空间对齐AI制图模型”串联在一起，实现了存量更新规划“体检-协商-生成-校验”的全数字化端到端闭环决策支持；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）基于地理红线刚性对齐的AIGC制图管线：设计「矢量-光栅-ControlNet」多通道控制机制，将GIS矢量红线转为光栅掩膜作为ControlNet条件约束，强制Stable Diffusion生成的规划图纸在建筑红线、道路中线及用地分区上与GIS矢量数据实现像素级空间位置对齐，从根本上消除了普通AIGC绘图的地理空间幻觉，满足法定规划的物理精度要求；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）大模型动态图例说明合成技术（Data-to-Text）：图册编译阶段直接读取底层的实测GIS空间指标，调用大模型生成量化设计说明文字，在Pillow画布中动态排版渲染，实现了“规划图面-量化指标-设计说明”的绝对绑定，消除手工修改滞后性；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（4）数据与业务逻辑的解耦架构：系统采用参数化配置与模块化设计。若将系统移植至其他城市的更新项目，仅需在 `data/` 目录下替换相应的GeoJSON和CSV文件，核心的诊断、博弈和制图引擎无需做出任何修改，具备极强的复用性与推广价值。", False, False, False)
    ])
replace_section_content_before(doc, "1．模型设计的特点", "2．应用方向或应用前景", populate_c6_s1)
# 第六章第二节
def populate_c6_s2(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本系统围绕高质量城市治理和存量城市更新，具备广泛的工程应用前景与落地推广价值：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）国家及地方城市体检与专项更新评估：系统可直接对接到住建部门的数字城市底座上，作为城市年度健康体检评估的量化测度工具，自动诊断并输出体检报告，定位更新重点；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）基层自然资源规划部门与街道决策辅助平台：可在规划听证会或街道更新协商会议上作为多方利益主体的数字推演工作坊，实时输入不同方案并自动刷新博弈满意度得分与合规警告，加速规划共识达成；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）高等院校规划与建筑设计教学示范：本系统可作为城乡规划及城市设计课程的数字化实验平台，引导学生直观理解多源大数据诊断、空间句法分析、智能体决策等数字规划的核心方法。", False, False, False)
    ])
# 参考文献
def populate_references(p_next):
    p_h = p_next.insert_paragraph_before(style="Heading 1")
    run_h = p_h.add_run("参考文献")
    set_font_run_only(run_h, font_name="Times New Roman", east_asia_font="黑体")
    run_h.bold = True
    p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    refs = [
        "[1] 丁梦月.基于计算机视觉技术的城市街道步行空间人群行为原型研究[D].南京:东南大学,2021.",
        "[2] 尧馨雅.基于可解释深度学习的街道风貌基因图谱识别研究[D].杭州:浙江大学,2022.",
        "[3] 赵卉.历史肌理延续的数字化城市设计方法研究——以江苏同里古镇为例[D].南京:东南大学,2021.",
        "[4] 卢文正.社会空间理论视角的社区更新[D].哈尔滨:哈尔滨工业大学,2020.",
        "[5] 张峰.智慧城市空间信息资源规划的模型和实现方法研究[D].武汉:武汉大学,2005.",
        "[6] 张国政.数字孪生技术提升城市韧性路径研究[D].上海:华东政法大学,2023.",
        "[7] 梁汉雄.基于街景图片与深度学习的旧工业园区改造与可步行性要素研究[D].广州:华南理工大学,2022.",
        "[8] 方可.探索北京旧城居住区有机更新的适宜途径[D].北京:清华大学,2000."
    ]
    for ref in refs:
        p_ref = p_next.insert_paragraph_before(style="Normal")
        run_ref = p_ref.add_run(ref)
        set_font_run_only(run_ref)
idx_other = find_paragraph_by_text(doc, "其他格式要求！！！")
if idx_other != -1:
    p_next_other = doc.paragraphs[idx_other]
    populate_references(p_next_other)
    replace_section_content_before(doc, "2．应用方向或应用前景", "参考文献", populate_c6_s2)
# 删除多余的页尾格式指南说明文字（保留分节符）
idx_other = find_paragraph_by_text(doc, "其他格式要求！！！")
if idx_other != -1:
    p_to_del_footer = [doc.paragraphs[i] for i in range(idx_other, len(doc.paragraphs))]
    for p in p_to_del_footer:
        if len(p._element.xpath('.//w:sectPr')) > 0:
            p.text = "" # 保留分节符
        else:
            p._element.getparent().remove(p._element)
# ── 4. 全局替换正文字体（排除封面页，保留原有的字号、粗细等格式） ──
print("正文及表格构建完成，执行全局字体替换...")
start_para_idx = 0
for idx, p in enumerate(doc.paragraphs):
    if p.text.strip() == "一、研究问题":
        start_para_idx = idx
        break
for idx, p in enumerate(doc.paragraphs):
    if idx < start_para_idx:
        # 封面和目录段落保持原样，绝不改动任何字体和字号
        continue
    style_name = p.style.name
    text = p.text.strip()
    if not text:
        continue
    # 标题、图表题注级段落设置 HeiTi + Times New Roman
    if style_name.startswith("Heading") or text.startswith(("一、", "二、", "三、", "四、", "五、", "六、", "参考文献", "图", "表")):
        for run in p.runs:
            set_font_run_only(run, font_name="Times New Roman", east_asia_font="黑体")
    # 普通正文段落设置 SongTi + Times New Roman
    else:
        for run in p.runs:
            set_font_run_only(run, font_name="Times New Roman", east_asia_font="宋体")
# 表格文本统一修改字体，绝不修改边框、底纹等格式
for idx, tbl in enumerate(doc.tables):
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font_run_only(run, font_name="Times New Roman", east_asia_font="宋体")
# ── 保存输出 ──
print(f"保存文档到: {OUTPUT_PATH}")
doc.save(str(OUTPUT_PATH))
print("✓ 完成！成功实现高保真度编译，未删除任何分节符且完美保留了封面字号样式。")


