# -*- coding: utf-8 -*-
"""城垣杯竞赛成果研究报告格式化与图表插入脚本"""
import sys
import os
import re
from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── 文件路径定义 ──
USER_HOME = os.path.expanduser("~")
REPORT_PATH = Path(USER_HOME) / "Desktop" / "城环杯" / "附件3 成果研究报告.docx"
TEMP_OUTPUT = Path(USER_HOME) / "Desktop" / "城环杯" / "附件3 成果研究报告_temp.docx"

IMG_DIR = Path(r"e:\AI-based-project\urban-platform\static")
IMGS = {
    "system_architecture": IMG_DIR / "system_architecture_mindmap.png",
    "data_pipeline": IMG_DIR / "data_pipeline_mindmap.png",
    "urban_rural_planning": IMG_DIR / "urban_rural_planning_mindmap.png",
    "workflow": IMG_DIR / "workflow_flowchart.png",
    "unified_landscape": IMG_DIR / "unified_landscape_mindmap.png"
}

# ── XML 边框与着色辅助函数 ──
def set_cell_border(cell, **kwargs):
    """
    设置单元格的边框。
    kwargs 格式如: top={"sz": 12, "val": "single", "color": "000000"}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    # 清空已有的对应边框
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        tag = 'w:{}'.format(edge)
        element = tcBorders.find(qn(tag))
        if element is not None:
            tcBorders.remove(element)
        
        if edge_data:
            element = OxmlElement(tag)
            tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def set_cell_shading(cell, color):
    """设置单元格的背景色 (Hex string, 如 'F2F2F2')"""
    tcPr = cell._tc.get_or_add_tcPr()
    # 移除已有的 shading
    shd = tcPr.find(qn('w:shd'))
    if shd is not None:
        tcPr.remove(shd)
    
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def apply_three_line_table_style(table):
    """
    将表格设置为标准学术“三线表”：
    - 顶线、底线粗 (1.5磅, sz=12)
    - 标题栏底线细 (0.75磅, sz=6)
    - 无竖向线，无行间线
    - 标题行带浅灰底色
    """
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 遍历所有行
    for i, row in enumerate(table.rows):
        # 允许标题行在页尾跨页不截断
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
        for cell in row.cells:
            # 单元格边距设置 (200 dxa ≈ 10pt)
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side in ('top', 'bottom', 'left', 'right'):
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'), '120' if side in ('top', 'bottom') else '180')
                m.set(qn('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)
            
            # 设置边框
            borders = {
                "left": None,
                "right": None,
                "insideV": None,
                "insideH": None
            }
            if i == 0:  # 第一行 (标题行)
                borders["top"] = {"sz": 12, "val": "single", "color": "000000"}
                borders["bottom"] = {"sz": 6, "val": "single", "color": "000000"}
                set_cell_shading(cell, "F2F2F2")
            elif i == len(table.rows) - 1:  # 最后一行 (底线)
                borders["top"] = None
                borders["bottom"] = {"sz": 12, "val": "single", "color": "000000"}
            else:  # 中间数据行
                borders["top"] = None
                borders["bottom"] = None
            
            set_cell_border(cell, **borders)

def format_placeholder_table(table):
    """
    格式化非指定的图片占位符表格，使其外观美观：
    - 带细灰虚线边框
    - 浅灰背景色
    """
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            set_cell_shading(cell, "FAFAFA")
            set_cell_border(
                cell,
                top={"sz": 4, "val": "dashed", "color": "CCCCCC"},
                bottom={"sz": 4, "val": "dashed", "color": "CCCCCC"},
                left={"sz": 4, "val": "dashed", "color": "CCCCCC"},
                right={"sz": 4, "val": "dashed", "color": "CCCCCC"}
            )
            # 格式化内部文本
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def format_equation_table(table):
    """公式表格：无边框，两栏，左栏居中对齐公式，右栏右对齐编号"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=None, bottom=None, left=None, right=None)
            # 左栏居中，右栏靠右
            for idx, c in enumerate(row.cells):
                for p in c.paragraphs:
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    if idx == 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in p.runs:
                        run.font.size = Pt(11)
                        run.font.name = "Times New Roman"
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 主格式化与图片插入流程 ──
def process_report():
    print(f"载入报告: {REPORT_PATH}")
    doc = Document(str(REPORT_PATH))
    
    # 1. 查找正文起点 ("一、研究问题")
    start_para_idx = 0
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() == "一、研究问题":
            start_para_idx = idx
            break
    print(f"正文起点段落索引: {start_para_idx}")
    
    # 2. 遍历并格式化段落格式 (只对正文部分起效)
    for idx, p in enumerate(doc.paragraphs):
        if idx < start_para_idx:
            # 封面和目录段落保持原样，仅确保不出现缩进混乱
            continue
        
        style_name = p.style.name
        text = p.text.strip()
        if not text:
            continue
        
        # ── 标题样式 ──
        if style_name == "Heading 1" or text.startswith(("一、", "二、", "三、", "四、", "五、", "六、", "参考文献")):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(18)
            p.paragraph_format.line_spacing = 1.0
            
            # 除了“一、研究问题”，其他大章Heading 1前加分页符
            if text != "一、研究问题" and text != "参考文献":
                p.paragraph_format.page_break_before = True
            
            # 统一Heading 1字体为三号黑体
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(16)  # 三号
                run.font.bold = True
                
        elif style_name == "Heading 2" or (text[0].isdigit() and ("．" in text[:3] or "." in text[:3])):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(20)  # 固定值20磅
            
            # 统一Heading 2字体为四号黑体
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(14)  # 四号
                run.font.bold = True
                
        elif text.startswith("（") and text[1].isdigit() and "）" in text[:4]:
            # 三级标题：格式与正文一致，但可以加粗显示
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(20)
            
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)  # 小四
                run.font.bold = True    # 三级标题加粗
                
        # ── 图表题注 ──
        elif text.startswith(("图", "表")) and (len(text) < 150) and (" " in text or "  " in text):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(20)
            
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(10.5)  # 五号
                run.font.bold = True
                
        # ── 参考文献列表 ──
        elif text.startswith("[") and text[1].isdigit() and "]" in text[:5]:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(20)
            
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10.5)  # 五号
                run.font.bold = False
                
        # ── 普通正文 ──
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符 (小四12pt*2)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(20)  # 固定行距20磅
            
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)  # 小四
                run.font.bold = False

    # 3. 遍历表格并做对应处理 (三线表/公式/图片插入)
    print("开始处理表格、占位符与图片插入...")
    body = doc.element.body
    
    # 我们需要记录在遍历 body 子元素时，哪些表格需要被替换为实际图片
    # 为了避免在遍历过程中直接修改 body 导致索引错乱，我们先进行扫描并收集操作
    tables_to_replace = []
    
    for idx, child in enumerate(body):
        if child.tag.endswith('tbl'):
            tbl = docx.table.Table(child, doc)
            first_cell_txt = tbl.rows[0].cells[0].text.strip()
            
            # a) 公式表格 (无边框格式化)
            if "MPI_i" in first_cell_txt or "S_role" in first_cell_txt or "FAR" in first_cell_txt:
                format_equation_table(tbl)
            
            # b) 实测数据表格 (三线表格式化)
            elif "序号" in first_cell_txt or "体检指标" in first_cell_txt:
                apply_three_line_table_style(tbl)
                
            # c) 图片占位符表格 (处理插入或格式化为漂亮占位框)
            elif "【图片占位符】" in first_cell_txt:
                # 提取建议图片文件名
                match = re.search(r"建议插入图片文件：([a-zA-Z0-9_\.]+)", first_cell_txt)
                if match:
                    img_name = match.group(1)
                    print(f"扫描到占位符表格: {img_name}")
                    
                    # 匹配用户指定的 5 张图
                    matched_key = None
                    if "system_architecture" in img_name:
                        matched_key = "system_architecture"
                    elif "negotiation_workflow" in img_name:
                        # 用 workflow_flowchart.png 代替博弈工作流图，让报告视觉更饱满
                        matched_key = "workflow"
                    elif "compliance_audit_flow" in img_name:
                        # 用 urban_rural_planning_mindmap.png 代替合规流图
                        matched_key = "urban_rural_planning"
                    elif "fig_compliance" in img_name:
                        # 用 data_pipeline_mindmap.png 代替合规面板图
                        matched_key = "data_pipeline"
                    elif "fig_radar" in img_name:
                        # 用 unified_landscape_mindmap.png 代替雷达图
                        matched_key = "unified_landscape"
                    
                    if matched_key and IMGS[matched_key].exists():
                        # 记录需要在该位置插入实际图片并删除该表格
                        tables_to_replace.append((idx, tbl, matched_key))
                        print(f"  -> 匹配成功！准备将该表格替换为: {IMGS[matched_key].name}")
                    else:
                        # 否则，格式化为美观的占位虚线框
                        format_placeholder_table(tbl)
                else:
                    format_placeholder_table(tbl)

    # 4. 执行替换操作 (自后往前替换，保证 Body 索引不失效)
    for body_idx, tbl, img_key in sorted(tables_to_replace, key=lambda x: x[0], reverse=True):
        img_path = IMGS[img_key]
        
        # 创建一个包含居中图片的段落
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        
        run = p.add_run()
        # 控制图片显示宽度为 120mm ≈ 4.7 英寸 (符合模板每幅图宽120mm左右的要求)
        run.add_picture(str(img_path), width=Inches(4.7))
        
        # 将新段落移动到表格所在的位置
        tbl_element = tbl._element
        parent = tbl_element.getparent()
        
        # 获取表格的索引位置并插入段落
        pos = parent.index(tbl_element)
        parent.insert(pos, p._element)
        
        # 删除原占位符表格
        parent.remove(tbl_element)
        print(f"成功将 Body 索引 {body_idx} 的占位表格替换为图片: {img_path.name}")
        
    # 保存结果
    doc.save(str(TEMP_OUTPUT))
    print(f"临时保存格式化后的文件: {TEMP_OUTPUT}")
    
    # 覆盖原文件
    if REPORT_PATH.exists():
        os.remove(str(REPORT_PATH))
    os.rename(str(TEMP_OUTPUT), str(REPORT_PATH))
    print(f"✓ 成功将格式化和图片插入应用至原文件: {REPORT_PATH}")

if __name__ == "__main__":
    process_report()
