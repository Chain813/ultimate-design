import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import base64

def generate_official_word_doc(title, text_content, aigc_history=None):
    """
    使用 python-docx 纯代码生成带红头的官方公文 Word 格式，
    支持插入 Markdown 转换来的纯文本和 AIGC 推演生成的 Base64 图像。
    """
    doc = docx.Document()
    
    # 1. 渲染大红头
    red_head = doc.add_paragraph()
    red_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh_run = red_head.add_run("长春市历史文化街区微更新规划专项报告")
    rh_run.font.size = Pt(22)
    rh_run.font.name = '宋体'
    rh_run.font.color.rgb = RGBColor(255, 0, 0)
    
    # 红线分隔符
    red_line = doc.add_paragraph()
    red_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl_run = red_line.add_run("★")
    rl_run.font.size = Pt(14)
    rl_run.font.color.rgb = RGBColor(255, 0, 0)
    
    # 2. 正文标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.space_before = Pt(20)
    title_p.space_after = Pt(20)
    t_run = title_p.add_run(title)
    t_run.font.size = Pt(18)
    t_run.font.bold = True

    # 3. 规划文本注入
    doc.add_heading("一、 微更新规划导则文本", level=1)
    
    # 简单的文本分段处理
    for paragraph in text_content.split('\n'):
        if paragraph.strip() == "":
            continue
        if paragraph.startswith('#'):
            # 处理加粗标题
            p = doc.add_paragraph()
            run = p.add_run(paragraph.replace('#', '').strip())
            run.font.bold = True
            run.font.size = Pt(14)
        elif paragraph.startswith('-') or paragraph.startswith('*'):
            # 处理列表
            p = doc.add_paragraph(paragraph.strip(), style='List Bullet')
        else:
            # 普通段落，首行缩进
            p = doc.add_paragraph(paragraph.strip())
            p.paragraph_format.first_line_indent = Pt(28) 

    # 4. AIGC 成果图集注入
    if aigc_history and len(aigc_history) > 0:
        doc.add_page_break()
        doc.add_heading("二、 重点地块 AIGC 视觉推演图集", level=1)
        
        for idx, img_data in enumerate(aigc_history):
            try:
                doc.add_heading(f"图景 {idx+1}: {img_data.get('plot', '未命名地块')}", level=2)
                p_desc = doc.add_paragraph()
                p_desc.add_run(f"采用策略：{img_data.get('strategy', '未定')}\n").font.bold = True
                p_desc.add_run(f"核心咒语记录：{img_data.get('prompt_excerpt', '')}")
                
                # 解码 Base64 并插入
                if 'thumb_b64' in img_data:
                    img_bytes = base64.b64decode(img_data['thumb_b64'])
                    img_io = BytesIO(img_bytes)
                    
                    pic_p = doc.add_paragraph()
                    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pic_p.add_run().add_picture(img_io, width=Inches(5.0))
            except Exception as e:
                doc.add_paragraph(f"[图片渲染失败: {e}]")

    # 5. 结尾声明
    doc.add_page_break()
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_p.add_run("本红头文件由 ultimateDESIGN 数字孪生平台自动排版生成。").font.color.rgb = RGBColor(128, 128, 128)
    
    # 存入缓存
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f
