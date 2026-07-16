"""项目设计报告 生成引擎

核心功能:
1. REPORT_CHAPTERS —— 27 个小节的完整定义（章节号、标题、字数、数据源、生成策略）
2. build_document_context() —— 从 stage_bus + spatial_data 聚合所有可用数据
3. generate_all_chapters() —— 逐章节调用 deepseek-v4-pro 生成正文
4. assemble_report_docx() —— 严格按模板格式组装 .docx

Usage:
    from src.engines.document_composer import generate_all_chapters, assemble_report_docx
    chapters = generate_all_chapters(student_info, progress_callback)
    docx_bytes = assemble_report_docx(chapters, student_info)
"""

from __future__ import annotations

import io
import logging
import re
from collections.abc import Callable
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger("ultimateDESIGN")


# ═══════════════════════════════════════════════════════════════
# 章节定义
# ═══════════════════════════════════════════════════════════════

@dataclass
class ReportSection:
    """项目设计报告 单个小节的完整定义"""
    section_id: str          # e.g. "1.1"
    title: str               # 中文标题
    word_count: int          # 目标字数
    chapter: int             # 所属章号
    strategy: str            # "rewrite" | "generate"
    data_sources: list[str] = field(default_factory=list)  # stage_bus keys / spatial hint
    description: str = ""    # 本节内容简述


REPORT_CHAPTERS: list[ReportSection] = [
    # ═══ 第 1 章：项目背景与概况 ═══
    ReportSection(
        section_id="1.1", title="项目背景", word_count=500, chapter=1,
        strategy="rewrite",
        data_sources=["design_brief", "diagnosis_report"],
        description="项目缘起、区域发展背景、城市更新政策背景、研究必要性",
    ),
    ReportSection(
        section_id="1.2", title="目标与任务", word_count=500, chapter=1,
        strategy="rewrite",
        data_sources=["design_brief", "design_concept", "case_benchmark"],
        description="规划设计目标、研究任务、技术路线概述",
    ),

    # ═══ 第 2 章：现状调查与分析 (每节 ~50 字概要) ═══
    ReportSection(
        section_id="2.1", title="区位及范围", word_count=50, chapter=2,
        strategy="generate",
        data_sources=["spatial_context", "boundary"],
        description="研究范围的地理位置、四至边界、面积",
    ),
    ReportSection(
        section_id="2.2", title="用地现状分析", word_count=50, chapter=2,
        strategy="generate",
        data_sources=["landuse_summary", "spatial_context"],
        description="现状用地分类与占比特征",
    ),
    ReportSection(
        section_id="2.3", title="交通现状分析", word_count=50, chapter=2,
        strategy="generate",
        data_sources=["traffic_summary", "spatial_context"],
        description="现状路网结构与交通特征",
    ),
    ReportSection(
        section_id="2.4", title="建筑现状分析", word_count=50, chapter=2,
        strategy="generate",
        data_sources=["building_summary", "spatial_context"],
        description="现状建筑高度、质量与分布特征",
    ),
    ReportSection(
        section_id="2.5", title="景观现状分析", word_count=50, chapter=2,
        strategy="generate",
        data_sources=["gvi_summary", "spatial_context"],
        description="现状绿化率、绿视率与景观品质",
    ),
    ReportSection(
        section_id="2.6", title="文化资源分析", word_count=50, chapter=2,
        strategy="generate",
        data_sources=["p04_cultural_analysis", "spatial_context"],
        description="历史建筑、工业遗产与文化资源分布",
    ),
    ReportSection(
        section_id="2.7", title="产业业态分析", word_count=50, chapter=2,
        strategy="generate",
        data_sources=["p04_industry_analysis", "poi_summary"],
        description="现状产业业态与商业分布特征",
    ),
    ReportSection(
        section_id="2.8", title="人群需求分析", word_count=50, chapter=2,
        strategy="generate",
        data_sources=["p04_population_analysis", "spatial_context"],
        description="现状人口特征与社区需求",
    ),

    # ═══ 第 3 章：设计理念与构思 (每节 ~200 字) ═══
    ReportSection(
        section_id="3.1", title="设计依据", word_count=200, chapter=3,
        strategy="rewrite",
        data_sources=["p07_design_basis", "design_concept", "case_benchmark", "rag_policy", "spatial_context"],
        description="上位规划、政策法规、案例分析依据",
    ),
    ReportSection(
        section_id="3.2", title="设计原则", word_count=200, chapter=3,
        strategy="rewrite",
        data_sources=["p07_design_principles", "design_concept", "strategy_matrix", "diagnosis_report"],
        description="规划设计基本原则（保护优先、有机更新、多方参与等）",
    ),
    ReportSection(
        section_id="3.3", title="设计目标", word_count=200, chapter=3,
        strategy="rewrite",
        data_sources=["design_concept", "design_brief", "diagnosis_report", "spatial_context"],
        description="量化设计目标（容积率、绿地率、限高等）",
    ),
    ReportSection(
        section_id="3.4", title="设计定位", word_count=200, chapter=3,
        strategy="rewrite",
        data_sources=["p07_design_positioning", "design_concept", "strategy_matrix", "spatial_structure"],
        description="项目功能定位与空间形象定位",
    ),
    ReportSection(
        section_id="3.5", title="设计策略", word_count=200, chapter=3,
        strategy="rewrite",
        data_sources=["strategy_matrix", "negotiation_result"],
        description="问题-目标-策略对应体系",
    ),

    # ═══ 第 4 章：总体方案设计 (每节 ~200 字) ═══
    ReportSection(
        section_id="4.1", title="总图说明", word_count=200, chapter=4,
        strategy="rewrite",
        data_sources=["spatial_structure", "landuse_sandbox"],
        description="总体规划布局说明",
    ),
    ReportSection(
        section_id="4.2", title="用地结构规划", word_count=200, chapter=4,
        strategy="rewrite",
        data_sources=["landuse_sandbox", "spatial_structure"],
        description="用地分类与结构优化",
    ),
    ReportSection(
        section_id="4.3", title="开发强度规划", word_count=200, chapter=4,
        strategy="rewrite",
        data_sources=["landuse_sandbox", "building_form", "spatial_structure"],
        description="容积率、建筑密度、高度分区",
    ),
    ReportSection(
        section_id="4.4", title="交通组织规划", word_count=200, chapter=4,
        strategy="rewrite",
        data_sources=["traffic_system", "spatial_structure"],
        description="道路系统、公共交通、慢行网络",
    ),
    ReportSection(
        section_id="4.5", title="空间布局规划", word_count=200, chapter=4,
        strategy="rewrite",
        data_sources=["spatial_structure", "public_space"],
        description="空间结构、轴线、节点体系",
    ),
    ReportSection(
        section_id="4.6", title="产业业态规划", word_count=200, chapter=4,
        strategy="generate",
        data_sources=["p09_industry_planning", "strategy_matrix", "spatial_structure"],
        description="产业布局与业态引导",
    ),
    ReportSection(
        section_id="4.7", title="景观系统规划", word_count=200, chapter=4,
        strategy="rewrite",
        data_sources=["landscape_style", "public_space", "gvi_summary"],
        description="绿地系统、景观结构、风貌分区",
    ),
    ReportSection(
        section_id="4.8", title="城市设计导则", word_count=200, chapter=4,
        strategy="rewrite",
        data_sources=["design_guideline", "design_brief"],
        description="城市设计管控导则概要",
    ),

    # ═══ 第 5 章：重点地块设计 ═══
    ReportSection(
        section_id="5.1", title="特色专项研究", word_count=200, chapter=5,
        strategy="generate",
        data_sources=["p10_specialized_study", "design_brief", "diagnosis_report"],
        description="项目特色专项（如工业遗产活化、数字孪生辅助设计等）",
    ),
    ReportSection(
        section_id="5.2", title="地块选择与依据", word_count=800, chapter=5,
        strategy="rewrite",
        data_sources=["mpi_ranking", "top_plot", "radar_data", "plot_designs", "plot_metrics", "spatial_context"],
        description="重点地块的选择标准、依据与概况",
    ),
    ReportSection(
        section_id="5.3", title="地块设计说明", word_count=800, chapter=5,
        strategy="rewrite",
        data_sources=["plot_designs", "plot_metrics", "plot_personas", "spatial_context", "landuse_sandbox"],
        description="重点地块详细设计方案",
    ),
]


# ═══════════════════════════════════════════════════════════════
# 数据聚合
# ═══════════════════════════════════════════════════════════════

def _safe_str(val, max_chars: int = 3000) -> str:
    """安全转字符串并截断"""
    if val is None:
        return ""
    s = str(val)
    if len(s) > max_chars:
        s = s[:max_chars] + "...(已截断)"
    return s


def build_document_context() -> dict[str, str]:
    """从 stage_bus + spatial_data 聚合所有可用数据源。

    Returns:
        dict: {source_name: text_content} 供 LLM prompt 注入使用
    """
    ctx: dict[str, str] = {}

    # ── 从 DesignContext 提取现有 stage 数据 ──
    try:
        from src.workflow.design_context import build_design_context as _bdc
        dc = _bdc()

        ctx["diagnosis_report"] = _safe_str(dc.diagnosis_report, 2000)
        ctx["mpi_ranking"] = _safe_str(dc.mpi_ranking, 1500)
        ctx["top_plot"] = _safe_str(dc.top_plot, 500)
        ctx["top_score"] = _safe_str(dc.top_score, 200)
        ctx["radar_data"] = _safe_str(dc.radar_data, 1000)
        ctx["design_concept"] = _safe_str(dc.design_concept, 2000)
        ctx["case_benchmark"] = _safe_str(dc.case_benchmark, 1500)
        ctx["strategy_matrix"] = _safe_str(dc.strategy_matrix, 1500)
        ctx["negotiation_result"] = _safe_str(dc.negotiation_result, 1500)
        ctx["spatial_structure"] = _safe_str(dc.spatial_structure, 2000)
        ctx["landuse_sandbox"] = _safe_str(dc.landuse_sandbox, 1500)
        ctx["traffic_system"] = _safe_str(dc.traffic_system, 1500)
        ctx["public_space"] = _safe_str(dc.public_space, 1500)
        ctx["building_form"] = _safe_str(dc.building_form, 1500)
        ctx["landscape_style"] = _safe_str(dc.landscape_style, 1500)
        ctx["design_guideline"] = _safe_str(dc.design_guideline, 2000)
        ctx["design_brief"] = _safe_str(dc.design_brief, 2000)

        # 地块级数据
        if dc.plot_designs:
            ctx["plot_designs"] = "\n".join(f"【{k}】{_safe_str(v, 1000)}" for k, v in dc.plot_designs.items())
        if dc.plot_metrics:
            ctx["plot_metrics"] = "\n".join(f"【{k}】{_safe_str(v, 800)}" for k, v in dc.plot_metrics.items())
        if dc.plot_personas:
            ctx["plot_personas"] = "\n".join(f"【{k}】{_safe_str(v, 800)}" for k, v in dc.plot_personas.items())
        if not dc.plot_designs:
            ctx["plot_designs"] = ""
        if not dc.plot_metrics:
            ctx["plot_metrics"] = ""
        if not dc.plot_personas:
            ctx["plot_personas"] = ""
    except Exception as e:
        logger.warning(f"Failed to build DesignContext: {e}")

    # ── 空间数据 ──
    try:
        from src.engines.spatial_data_injector import (
            get_building_summary,
            get_full_spatial_context,
            get_gvi_summary,
            get_landuse_summary,
            get_poi_summary,
            get_traffic_summary,
        )
        ctx["spatial_context"] = _safe_str(get_full_spatial_context(), 3000)
        ctx["landuse_summary"] = _safe_str(get_landuse_summary(), 1000)
        ctx["poi_summary"] = _safe_str(get_poi_summary(), 1000)
        ctx["gvi_summary"] = _safe_str(get_gvi_summary(), 1000)
        ctx["building_summary"] = _safe_str(get_building_summary(), 1000)
        ctx["traffic_summary"] = _safe_str(get_traffic_summary(), 1000)
    except Exception as e:
        logger.warning(f"Failed to load spatial data: {e}")

    # ── RAG 政策依据 ──
    try:
        from src.engines.rag_engine import retrieve_rag_context
        chunks = retrieve_rag_context("长春市历史街区保护规划设计导则", top_k=5)
        if chunks:
            ctx["rag_policy"] = "\n\n".join(f"[{c[2]}]: {c[1]}" for c in chunks)
    except Exception:
        ctx["rag_policy"] = ""

    # ── 缺失章节补充数据 (从 stage_bus 直接读取) ──
    try:
        from src.workflow.stage_data_bus import load_stage_output
        from src.workflow.stage_keys import SK

        ctx["p04_cultural_analysis"] = _safe_str(load_stage_output("04", SK.CULTURAL_ANALYSIS, ""), 800)
        ctx["p04_industry_analysis"] = _safe_str(load_stage_output("04", SK.INDUSTRY_ANALYSIS, ""), 800)
        ctx["p04_population_analysis"] = _safe_str(load_stage_output("04", SK.POPULATION_ANALYSIS, ""), 800)

        ctx["p07_design_basis"] = _safe_str(load_stage_output("07", SK.DESIGN_BASIS, ""), 800)
        ctx["p07_design_principles"] = _safe_str(load_stage_output("07", SK.DESIGN_PRINCIPLES, ""), 800)
        ctx["p07_design_positioning"] = _safe_str(load_stage_output("07", SK.DESIGN_POSITIONING, ""), 800)

        ctx["p09_industry_planning"] = _safe_str(load_stage_output("09", SK.INDUSTRY_PLANNING, ""), 1000)

        ctx["p10_specialized_study"] = _safe_str(load_stage_output("10", SK.SPECIALIZED_STUDY, ""), 1000)
    except Exception as e:
        logger.warning(f"Failed to load supplementary stage data: {e}")

    return ctx


# ═══════════════════════════════════════════════════════════════
# LLM Prompt 构建
# ═══════════════════════════════════════════════════════════════

def get_document_system_prompt() -> str:
    from src.config import get_site_city, get_site_desc, get_site_district, get_site_name
    city = get_site_city()
    district = get_site_district()
    site_name = get_site_name()
    desc = get_site_desc()

    from src.config.site import get_project_info
    proj = get_project_info()
    proj_name = proj.get("name", "城市设计智能推演平台")
    proj_sub = proj.get("subtitle", f"——以{city}市{district}{site_name}为例")

    return f"""你是一个严格的数据到文本转换引擎。你的唯一任务是将提供的项目数据
整理成项目设计报告的正文段落。

项目名称：《{proj_name}
{proj_sub}》

研究范围：{city}市{district}{site_name}，{desc}

【核心约束 — 违反任何一条即为失败】
1. 你只能使用下方提供的源材料中的信息。源材料中出现的数值、地名、事实可以引用；源材料中没有的，一律不得出现。
2. 禁止角色扮演。你不是"学生""规划师""专家"，你是一个数据整理工具。
3. 禁止插入个人观察、个人经历、个人感受（如"笔者走访""实地调研发现""给人留下深刻印象"）。你不是人类，没有个人经历。
4. 禁止凭空生成案例、数据、地名、人物、法规条文编号。如果源材料没有提供，就留白。
5. 禁止添加任何评价性语言（如"具有重要价值""意义深远""值得关注"）。只陈述事实。
6. 禁止使用"首先/其次/最后""综上所述""值得注意的是"等模板化过渡词。
7. 禁止使用"赋能""织补""触媒""韧性"等 AI 高频词汇。

【写作格式】
1. 第三人称，客观陈述。使用规范的城乡规划术语。
2. 直接陈述结论，不描述分析过程。
3. 每段 2-5 句话，句长自然变化。
4. 第2章各小节为摘要式（约50字），仅给出核心特征。
5. 只输出正文段落，不输出章节标题。"""


def _resolve_sources(sec: ReportSection, ctx: dict[str, str]) -> str:
    """将 data_sources 列表解析为可注入 prompt 的文本块"""
    parts = []
    for src in sec.data_sources:
        val = ctx.get(src, "")
        if val:
            label = {
                "diagnosis_report": "前期诊断报告",
                "mpi_ranking": "MPI 地块排行",
                "top_plot": "重点地块",
                "top_score": "最高 MPI 得分",
                "radar_data": "地块雷达数据",
                "design_concept": "设计概念报告",
                "case_benchmark": "案例对标分析",
                "strategy_matrix": "策略矩阵",
                "negotiation_result": "多方协商结果",
                "spatial_structure": "空间结构推演",
                "landuse_sandbox": "用地沙盘模拟",
                "traffic_system": "交通系统设计",
                "public_space": "公共空间设计",
                "building_form": "建筑形态设计",
                "landscape_style": "风貌景观设计",
                "design_guideline": "城市设计导则",
                "design_brief": "设计纲要",
                "plot_designs": "地块设计方案",
                "plot_metrics": "地块控规指标",
                "plot_personas": "地块人群画像",
                "spatial_context": "空间数据总览",
                "landuse_summary": "用地统计",
                "poi_summary": "POI 分布",
                "gvi_summary": "街景品质",
                "building_summary": "建筑统计",
                "traffic_summary": "交通统计",
                "rag_policy": "政策法规依据",
                "p04_cultural_analysis": "文化资源分析",
                "p04_industry_analysis": "产业业态分析",
                "p04_population_analysis": "人群需求分析",
                "p07_design_basis": "设计依据",
                "p07_design_principles": "设计原则",
                "p07_design_positioning": "设计定位",
                "p09_industry_planning": "产业业态规划",
                "p10_specialized_study": "特色专项研究",
            }.get(src, src)
            parts.append(f"【{label}】\n{val}")
    return "\n\n".join(parts)


def build_chapter_prompt(sec: ReportSection, ctx: dict[str, str]) -> str:
    """为单个小节构建 LLM prompt"""
    source_text = _resolve_sources(sec, ctx)
    has_sources = bool(source_text.strip())

    if not has_sources:
        # 无源数据时绝不编造，输出占位符提示用户先生成上游数据
        return f"""【章节】{sec.section_id} {sec.title}
【字数要求】约{sec.word_count}字

⚠️ 本节无可用源数据。你必须输出以下占位符文本（一字不差）：

[待生成] 本节依赖上游阶段数据，请先在对应阶段页面运行 AI 生成后再重新生成设计报告。依赖数据源：{', '.join(sec.data_sources) if sec.data_sources else '无特定数据源'}。

禁止在占位符之外输出任何其他内容。"""

    if sec.strategy == "rewrite":
        return f"""【章节】{sec.section_id} {sec.title}
【字数要求】约{sec.word_count}字
【内容说明】{sec.description}

请将以下源材料严格压缩改写为本节正文：

【源材料开始】
{source_text}
【源材料结束】

【改写铁律 — 逐一核对】
1. 约{sec.word_count}字，不得为凑字数而重复或添加内容
2. 所有数值（面积、百分比、指标、编号）必须与源材料完全一致，改一个数字即为错误
3. 所有地名必须与源材料完全一致，不得替换、泛化或具体化
4. 源材料中没有出现的事实、案例、数据、地名、人名、法规编号 → 不得出现
5. 源材料中已出现但你不确定含义的内容 → 原样保留，不要解释
6. 删除过程性描述（"通过分析""可以看出""值得注意"），直接陈述结论
7. 删除评价性语言（"具有重要意义""影响深远"），只留事实
8. 不要添加任何源材料之外的句子来"润色"或"丰富内容"
9. 源材料不足时，写短不写长，绝不填充编造
10. 只输出正文段落，不输出章节标题"""

    # 默认/生成策略
    return f"""【章节】{sec.section_id} {sec.title}
【字数要求】约{sec.word_count}字
【内容说明】{sec.description}

请基于以下项目源材料，生成本节的正文段落：

【源材料开始】
{source_text}
【源材料结束】

【生成铁律 — 逐一核对】
1. 约{sec.word_count}字，结构紧凑，陈述客观，术语规范。
2. 所有数值（面积、百分比、指标、编号）和地名必须与源材料完全一致，严禁任何形式的编造或泛化。
3. 严格基于数据事实，不得填充任何源材料中未提及的案例、规划内容或未来设想。
4. 只输出正文段落，不输出章节标题。"""


# ═══════════════════════════════════════════════════════════════
# 章节生成编排
# ═══════════════════════════════════════════════════════════════

def generate_single_section(
    sec: ReportSection,
    ctx: dict[str, str],
    model: str = "deepseek-v4-pro",
) -> str:
    """生成单个小节的正文内容"""
    from src.engines.llm_engine import call_llm_engine

    prompt = build_chapter_prompt(sec, ctx)
    result = call_llm_engine(
        prompt=prompt,
        system_prompt=get_document_system_prompt(),
        model=model,
    )

    if not result or len(result) < 10:
        return f"[生成失败] 请检查 LLM API 连接后重试。章节: {sec.section_id} {sec.title}"

    return result.strip()


def generate_all_chapters(
    ctx: dict[str, str] | None = None,
    progress_callback: Callable[[int, int, str], None] | None = None,
    chunk_callback: Callable[[str, str, str], None] | None = None,
    model: str = "deepseek-v4-pro",
) -> dict[str, str]:
    """生成全部 27 个小节。

    Args:
        ctx: 预构建的数据上下文（为 None 时自动构建）
        progress_callback: (current, total, section_id) 进度回调
        chunk_callback: (section_id, title, content) 每节完成回调
        model: LLM 模型

    Returns:
        dict: {section_id: generated_text}
    """
    import streamlit as st

    if ctx is None:
        ctx = build_document_context()

    results: dict[str, str] = {}
    total = len(REPORT_CHAPTERS)

    for i, sec in enumerate(REPORT_CHAPTERS):
        if progress_callback:
            progress_callback(i, total, sec.section_id)

        try:
            text = generate_single_section(sec, ctx, model=model)
            results[sec.section_id] = text

            if chunk_callback:
                chunk_callback(sec.section_id, sec.title, text)

        except Exception as e:
            logger.error(f"Failed to generate section {sec.section_id}: {e}")
            results[sec.section_id] = f"[生成异常] {sec.section_id} {sec.title}: {e}"

    if progress_callback:
        progress_callback(total, total, "完成")

    return results


# ═══════════════════════════════════════════════════════════════
# Docx 格式化工具 (复用自 build_final_report_strict_v3.py)
# ═══════════════════════════════════════════════════════════════

def _set_font_run(run, font_name="Times New Roman", east_asia_font="宋体"):
    """设置 run 的中英文字体"""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), east_asia_font)


def _add_body_paragraph(doc, text: str, first_indent: bool = True):
    """添加正文段落（宋体 12pt, 1.5 倍行距, 首行缩进 2 字符）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_indent:
        pf.first_line_indent = Pt(24)  # 约 2 字符

    run = p.add_run(text)
    _set_font_run(run)
    run.font.size = Pt(12)
    return p


def _add_chapter_heading(doc, text: str):
    """添加章标题（黑体 16pt bold 居中，段前分页）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    # 模板要求：一级标题单倍行距，段前24磅，段后18磅
    pf.line_spacing = 1.0
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    pf.first_line_indent = Pt(0)

    run = p.add_run(text)
    _set_font_run(run, font_name="Times New Roman", east_asia_font="黑体")
    run.font.size = Pt(16)
    run.bold = True
    return p


def _add_section_heading(doc, text: str):
    """添加节标题（黑体 14pt bold, 左对齐，段前24磅，段后6磅）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(24)
    pf.space_after = Pt(6)
    pf.first_line_indent = Pt(0)

    run = p.add_run(text)
    _set_font_run(run, font_name="Times New Roman", east_asia_font="黑体")
    run.font.size = Pt(14)
    run.bold = True
    return p


def _add_centered_line(doc, text: str, font_size: int = 12, bold: bool = False,
                        east_asia: str = "宋体", space_after: int = 6):
    """添加居中行（用于封面）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(space_after)
    pf.first_line_indent = Pt(0)

    run = p.add_run(text)
    _set_font_run(run, east_asia_font=east_asia)
    run.font.size = Pt(font_size)
    if bold:
        run.bold = True
    return p


def _add_page_break(doc):
    """添加分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)


def _add_toc(doc):
    """插入目录域代码（需用户在 Word 中右键更新域）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(24)
    pf.first_line_indent = Pt(0)

    run = p.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar_begin)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-2" \\h \\z '
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run3._element.append(fldChar_separate)

    run4 = p.add_run('（请在 Word 中右键此处 → 更新域 → 更新整个目录）')
    _set_font_run(run4)
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    run5 = p.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run5._element.append(fldChar_end)


# ═══════════════════════════════════════════════════════════════
# Docx 组装
# ═══════════════════════════════════════════════════════════════

@dataclass
class AuthorInfo:
    """学生信息"""
    name: str = ""
    student_id: str = ""
    advisor: str = ""
    college: str = "建筑与规划学院"
    major: str = "城乡规划"
    date: str = "2026年6月"


def load_author_info_json() -> AuthorInfo:
    """从 config/student_info.json 加载学生配置，确保不硬编码隐私数据"""
    import json
    import os

    from src.config.runtime import resolve_path
    
    default_info = AuthorInfo()
    try:
        config_path = resolve_path("config/student_info.json")
        if config_path.exists():
            with open(config_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                return AuthorInfo(
                    name=data.get("name", ""),
                    student_id=data.get("student_id", ""),
                    advisor=data.get("advisor", ""),
                    college=data.get("college", default_info.college),
                    major=data.get("major", default_info.major),
                    date=data.get("date", default_info.date),
                )
    except Exception as e:
        logger.warning(f"Failed to load config/student_info.json: {e}")
    return default_info


def save_author_info_json(student: AuthorInfo) -> None:
    """将学生学籍配置写入本地 config/student_info.json (已被 gitignore 忽略)"""
    import json
    import os

    from src.config.runtime import resolve_path
    
    try:
        config_dir = resolve_path("config")
        config_dir.mkdir(exist_ok=True)
        config_path = config_dir / "student_info.json"
        
        data = {
            "name": student.name,
            "student_id": student.student_id,
            "advisor": student.advisor,
            "college": student.college,
            "major": student.major,
            "date": student.date,
        }
        with open(config_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
    except Exception as e:
        logger.warning(f"Failed to save config/student_info.json: {e}")


def scan_local_references() -> list[str]:
    """扫描本地参考文献文件夹中的所有 PDF 文件名"""
    import os
    from pathlib import Path

    from src.config.paths import ROOT_DIR, config

    # Try config first
    path_val = config.get("data", {}).get("references_dir")
    if path_val:
        path = Path(path_val)
        if not path.is_absolute():
            path = ROOT_DIR / path
        path = str(path)
    else:
        # Fall back to workspace default
        path = str(ROOT_DIR / "data" / "references")

    if os.path.exists(path):
        try:
            files = [f for f in os.listdir(path) if f.lower().endswith('.pdf')]
            files.sort()
            return files
        except Exception as e:
            logger.warning(f"Failed to scan local references: {e}")
            return []
    return []


def assemble_report_docx(
    chapters: dict[str, str],
    student: AuthorInfo | None = None,
    abstract_cn: str = "",
    abstract_en: str = "",
    keywords_cn: str = "",
    keywords_en: str = "",
    references: str = "",
    acknowledgments: str = "",
) -> io.BytesIO:
    """组装完整的项目设计报告 .docx

    Returns:
        BytesIO buffer 用于 st.download_button
    """
    if student is None:
        student = AuthorInfo()

    doc = Document()

    # ── 全局页面设置 ──
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ════════════════════════════════════
    # 中文摘要
    # ════════════════════════════════════
    _add_centered_line(doc, "摘  要", font_size=16, bold=True, east_asia="黑体", space_after=12)
    ab_cn = abstract_cn or _generate_abstract_from_chapters(chapters)
    _add_body_paragraph(doc, ab_cn, first_indent=True)
    _add_body_paragraph(doc, "", first_indent=False)

    kw_cn = keywords_cn or _extract_keywords_from_chapters(chapters)
    p_kw = doc.add_paragraph()
    p_kw.paragraph_format.first_line_indent = Pt(0)
    p_kw.paragraph_format.line_spacing = 1.5
    run_kw_label = p_kw.add_run("关键词：")
    _set_font_run(run_kw_label)
    run_kw_label.font.size = Pt(12)
    run_kw_label.bold = True
    run_kw_val = p_kw.add_run(kw_cn)
    _set_font_run(run_kw_val)
    run_kw_val.font.size = Pt(12)

    _add_page_break(doc)

    # ════════════════════════════════════
    # English Abstract
    # ════════════════════════════════════
    _add_centered_line(doc, "Abstract", font_size=16, bold=True, east_asia="Times New Roman", space_after=12)
    ab_en = abstract_en or _generate_english_abstract(chapters)
    _add_body_paragraph(doc, ab_en, first_indent=True)
    _add_body_paragraph(doc, "", first_indent=False)

    kw_en_text = keywords_en or _extract_english_keywords(chapters)
    p_kwe = doc.add_paragraph()
    p_kwe.paragraph_format.first_line_indent = Pt(0)
    p_kwe.paragraph_format.line_spacing = 1.5
    run_kwe_label = p_kwe.add_run("Keywords: ")
    _set_font_run(run_kwe_label, east_asia_font="Times New Roman")
    run_kwe_label.font.size = Pt(12)
    run_kwe_label.bold = True
    run_kwe_val = p_kwe.add_run(kw_en_text)
    _set_font_run(run_kwe_val, east_asia_font="Times New Roman")
    run_kwe_val.font.size = Pt(12)

    _add_page_break(doc)

    # ════════════════════════════════════
    # 目录
    # ════════════════════════════════════
    _add_centered_line(doc, "目  录", font_size=16, bold=True, east_asia="黑体", space_after=12)
    _add_toc(doc)

    # ── 分节：正文开始新节（用于页码控制）──
    # 前导页（封面/声明/摘要/目录）为第 0 节，正文为第 1 节
    body_section = doc.add_section()
    body_section.top_margin = Cm(2.54)
    body_section.bottom_margin = Cm(2.54)
    body_section.left_margin = Cm(3.18)
    body_section.right_margin = Cm(3.18)
    # 正文节页脚：居中阿拉伯数字页码，起始页码为 1
    body_footer = body_section.footer
    body_footer.is_linked_to_previous = False
    footer_para = body_footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 插入 PAGE 域代码
    footer_run = footer_para.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    footer_run._element.append(fld_begin)
    fld_code = footer_para.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_code._element.append(instr)
    fld_sep = footer_para.add_run()
    fld_sep_char = OxmlElement('w:fldChar')
    fld_sep_char.set(qn('w:fldCharType'), 'separate')
    fld_sep._element.append(fld_sep_char)
    fld_num = footer_para.add_run('1')
    _set_font_run(fld_num)
    fld_num.font.size = Pt(10.5)
    fld_end = footer_para.add_run()
    fld_end_char = OxmlElement('w:fldChar')
    fld_end_char.set(qn('w:fldCharType'), 'end')
    fld_end._element.append(fld_end_char)
    # 前导节：无页脚（封面等不显示页码）
    for i, sec in enumerate(doc.sections):
        if i == 0:
            sec.different_first_page_header_footer = False
            # 第 0 节页脚设为空
            if sec.footer.paragraphs:
                pass  # 默认空页脚，不添加内容

    # ════════════════════════════════════
    # 正文 5 章
    # ════════════════════════════════════
    chapter_names = {
        1: "第1章  项目背景与概况",
        2: "第2章  现状调查与分析",
        3: "第3章  设计理念与构思",
        4: "第4章  总体方案设计",
        5: "第5章  重点地块设计",
    }

    last_chapter = 0
    for sec in REPORT_CHAPTERS:
        text = chapters.get(sec.section_id, f"[未生成] {sec.title}")

        # 新章开始时添加章标题
        if sec.chapter != last_chapter:
            if last_chapter != 0 and sec.chapter <= 5:
                _add_page_break(doc)
            ch_name = chapter_names.get(sec.chapter, f"第{sec.chapter}章")
            _add_chapter_heading(doc, ch_name)
            last_chapter = sec.chapter

        # 节标题
        _add_section_heading(doc, f"{sec.section_id}  {sec.title}")

        # 正文段落：按换行拆分
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        for para_text in paragraphs:
            # 跳过看起来像标题的行
            if para_text.startswith('#') or para_text.startswith('【'):
                continue
            _add_body_paragraph(doc, para_text, first_indent=True)

    _add_page_break(doc)

    # ════════════════════════════════════
    # 参考文献
    # ════════════════════════════════════
    _add_centered_line(doc, "参考文献", font_size=16, bold=True, east_asia="黑体", space_after=12)

    refs_text = references or _generate_references_from_chapters(chapters)
    refs = [r.strip() for r in refs_text.split('\n') if r.strip() and r.strip().startswith('[')]
    if not refs:
        refs = [refs_text]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # 悬挂缩进
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = Pt(-24)
        run = p.add_run(ref)
        _set_font_run(run)
        run.font.size = Pt(10.5)  # 五号字

    _add_page_break(doc)

    # ════════════════════════════════════
    # 致谢
    # ════════════════════════════════════
    _add_centered_line(doc, "致  谢", font_size=16, bold=True, east_asia="黑体", space_after=12)

    ack = acknowledgments or _generate_acknowledgments(student)
    _add_body_paragraph(doc, ack, first_indent=True)

    # ════════════════════════════════════
    # 附录（精简）
    # ════════════════════════════════════
    _add_page_break(doc)
    _add_centered_line(doc, "附  录", font_size=16, bold=True, east_asia="黑体", space_after=12)
    _add_body_paragraph(doc,
        "附录一：项目技术路线图\n"
        "（详见成果展示页面——全阶段生成汇总报告）",
        first_indent=False,
    )
    _add_body_paragraph(doc,
        "附录二：重点更新单元 MPI 评价结果表\n"
        "（详见 Stage 05 MPI 更新潜力评估——导出 CSV）",
        first_indent=False,
    )

    # ── 输出到 BytesIO ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════════════
# 辅助生成函数 (摘要/关键词/致谢/参考文献)
# ═══════════════════════════════════════════════════════════════

def _all_chapters_text(chapters: dict[str, str], max_chars: int = 5000) -> str:
    """拼接所有章节文本"""
    parts = []
    total = 0
    for sec in REPORT_CHAPTERS:
        text = chapters.get(sec.section_id, "")
        if text:
            excerpt = text[:300]
            parts.append(f"{sec.section_id} {sec.title}: {excerpt}")
            total += len(excerpt)
            if total > max_chars:
                break
    return "\n".join(parts)


def _generate_abstract_from_chapters(chapters: dict[str, str]) -> str:
    """从章节内容生成中文摘要——聚焦LLM与AIGC在城乡规划中的应用"""
    from src.engines.llm_engine import call_llm_engine

    ctx = _all_chapters_text(chapters)
    prompt = f"""请撰写一份约 400 字的中文学术摘要，严格聚焦于"大语言模型（LLM）与 AIGC 在城乡规划设计中的应用"这一核心主题。

【论文题目】《基于大模型与多模态AI的城市更新空间设计智能推演系统——以长春市宽城区伪满皇宫周边街区为例》
【研究区域】长春市宽城区伪满皇宫周边街区，约 160 公顷

【摘要主题定位】
本设计的核心创新不在于对特定地块的规划方案本身，而在于**提出并验证了一套以大语言模型和生成式AI为核心的城市空间智能推演方法论**。摘要必须以此为纲。

【摘要结构要求】
1. 背景（1-2 句）：存量更新时代，传统城市规划依赖人工经验，效率低且难以量化多方诉求。大语言模型与 AIGC 技术的突破为城市设计提供了新的技术范式。
2. 目的（1 句）：本项目构建了"基于大模型与多模态AI的城市更新空间设计智能推演系统"，探索 LLM 驱动空间诊断、多智能体协商、AIGC 方案生成的技术路径。
3. 方法（3-4 句，重点展开）：
   - 基于 DeepSeek 大语言模型的多智能体协商框架：模拟居民、开发商、规划师三方角色，进行三轮动态博弈并达成策略共识；
   - AIGC 设计推演引擎：以 Stable Diffusion 为核心，结合 ControlNet GIS 空间约束（路网骨架、建筑肌理、红线边界），实现"数据→提示词→可控图像"的生成式设计管线；
   - LLM 驱动的空间诊断与报告生成：整合 POI、街景绿视率（GVI）、空间句法等多源数据，由 LLM 进行 MPI 更新潜力评价并自动撰写诊断报告；
   - RAG 增强的政策校验：以向量检索注入地方规划法规，约束 LLM 输出合规性。
4. 主要发现与结论（2-3 句，定性概括）：
   - 多智能体 LLM 协商能有效平衡保护与更新的多方利益冲突；
   - AIGC+ControlNet 可在严守空间边界的前提下产出具有设计参考价值的方案意象图；
   - LLM 驱动的全流程自动化管线将设计推演周期从周级压缩至分钟级。
5. 意义（1 句）：本系统为 AI 辅助城市设计提供了可复用的技术框架，验证了大模型与 AIGC 在城乡规划领域的工程可行性。

【重要约束】
- 以"LLM + AIGC + 城乡规划"为主线，技术方法描述应占摘要篇幅的 50% 以上
- 定性概括为主，全文最多 2 个具体数字
- 不得使用"赋能""织补""触媒""韧性"等 AI 高频词汇
- 不得使用"首先/其次/最后"组织段落

【章节内容参考】
{ctx}

只输出摘要正文，不包含标题。"""
    result = call_llm_engine(
        prompt=prompt,
        system_prompt="你是一个学术摘要撰写工具。严格基于提供的章节内容撰写摘要，摘要中所有技术方法的描述必须与正文实际使用的方法一致。不得编造正文中没有的技术方案、实验数据或结论。使用规范的学术语言，第三人称，客观简洁。",
        model="deepseek-v4-pro",
    )
    return result.strip() if result else "摘要生成失败，请检查 LLM API。"


def _extract_keywords_from_chapters(chapters: dict[str, str]) -> str:
    """从章节内容提取中文关键词——突出AI与AIGC主题"""
    from src.engines.llm_engine import call_llm_engine

    ctx = _all_chapters_text(chapters, 2000)
    prompt = f"""请根据以下项目设计内容，提取 3-5 个核心中文关键词（用分号分隔）。

项目题目：《基于大模型与多模态AI的城市更新空间设计智能推演系统——以长春市宽城区伪满皇宫周边街区为例》

关键词应优先覆盖以下核心维度：
- 大语言模型（LLM）与多智能体协商
- AIGC / 生成式AI辅助设计
- 城市更新与历史街区保护
- 空间数据分析（GIS、空间句法、绿视率）
- 数字孪生与智能推演

内容参考：
{ctx}

只输出关键词，用分号分隔。关键词必须包含至少1个AI技术类术语和至少1个城市规划类术语。"""
    result = call_llm_engine(
        prompt=prompt,
        system_prompt="你是学术关键词提取工具。严格基于内容提取3-5个核心关键词，不添加内容中未涉及的主题词。",
        model="deepseek-v4-pro",
    )
    return result.strip() if result else "城市更新；大语言模型；AIGC；多智能体协商；AI辅助设计"


def _generate_english_abstract(chapters: dict[str, str]) -> str:
    """生成英文摘要——聚焦LLM与AIGC在城乡规划中的应用"""
    from src.engines.llm_engine import call_llm_engine

    cn_abstract = _generate_abstract_from_chapters(chapters)
    prompt = f"""Translate the following Chinese academic abstract into English (250-350 words),
maintaining the STRONG focus on LLM (Large Language Model) and AIGC (AI-Generated Content) applications
in urban-rural planning.

Important:
- The paper title: "An Intelligent Urban Renewal Spatial Design Inference System Based on Large Language Models and Multimodal AI: A Case Study of the Area Surrounding the Puppet Emperor's Palace in Changchun"
- Study area: ~160 hectares around the Puppet Emperor's Palace (Weimanhuang Palace), Kuancheng District, Changchun
- Core innovation: LLM-driven multi-agent negotiation + AIGC (Stable Diffusion + ControlNet) design generation + LLM spatial diagnosis pipeline
- Preserve the qualitative, methodology-focused tone — do NOT add numbers not in the original
- Key technical terms: MPI (Micro-renewal Potential Index), GVI (Green View Index), RAG (Retrieval-Augmented Generation), ControlNet, space syntax
- The English abstract should read as a methodology paper abstract in an urban planning / computational design journal

Chinese abstract to translate:
{cn_abstract}

Output only the English abstract text. No title."""
    result = call_llm_engine(
        prompt=prompt,
        system_prompt="You are an academic translation tool. Translate the Chinese abstract precisely into natural academic English. Do not add content not present in the original, and do not fabricate data, statistics, or findings.",
        model="deepseek-v4-pro",
    )
    return result.strip() if result else "Abstract generation failed."


def _extract_english_keywords(chapters: dict[str, str]) -> str:
    """提取英文关键词"""
    from src.engines.llm_engine import call_llm_engine

    ctx = _all_chapters_text(chapters, 2000)
    prompt = f"""Based on the following graduation project content, extract 3-5 English keywords (separated by semicolons).

Content:
{ctx}

Output only keywords, e.g.: Urban Renewal; Historic District; Digital Twin; AIGC; Micro-Renewal"""
    result = call_llm_engine(
        prompt=prompt,
        system_prompt="You are an academic keyword extraction tool. Extract keywords strictly based on the provided content. Do not add topics not covered in the text.",
        model="deepseek-v4-pro",
    )
    return result.strip() if result else "Urban Renewal; Historic District; Digital Twin; AI-Assisted Design"


def _generate_references_from_chapters(chapters: dict[str, str]) -> str:
    """生成 GB/T 7714-2015 格式参考文献"""
    from src.engines.llm_engine import call_llm_engine

    ctx = _all_chapters_text(chapters, 3000)
    local_refs = scan_local_references()
    local_refs_str = "\n".join(f"- {f}" for f in local_refs) if local_refs else "无"

    prompt = f"""请根据以下项目设计内容，以及学生本地阅读的真实文献，生成至少 30 篇参考文献列表，严格遵循 GB/T 7714-2015 格式。

【学生本地阅读的真实文献】（你必须首先为以下真实文献文件名生成标准的 GB/T 7714-2015 学术文献条目，并排在整个参考文献的最前面）：
{local_refs_str}

【其他相关文献生成要求】：
请补足其余的参考文献，使总数达到至少 30 篇以上，且英文文献占比不低于 30%（即英文文献不少于 9 篇），近 5 年（2021-2026）文献不少于 10 篇。主要涵盖期刊论文[J]、专著[M]、学位论文[D]等类型。
请按编号格式 [1] [2] ... [30] 顺序排列，合并输出。

主题相关领域：
- 城市更新与历史街区保护
- 数字孪生与智慧城市
- AIGC 与人工智能辅助设计
- 空间句法与城市形态
- 街景图像与计算机视觉
- 多主体利益协商与公众参与

内容参考：
{ctx}

请直接输出参考文献列表，每行一条，无需其他解释。"""
    result = call_llm_engine(
        prompt=prompt,
        system_prompt="你是学术参考文献生成工具。生成与论文内容主题相关的参考文献，格式遵循 GB/T 7714-2015。优先输出经典和公认的学术文献，不要编造不存在的作者、期刊或DOI。",
        model="deepseek-v4-pro",
    )
    return result.strip() if result else "[1] 参考文献生成失败，请检查 LLM API。"


def _generate_acknowledgments(student: AuthorInfo) -> str:
    """生成致谢"""
    from src.config.site import get_institution_info, get_project_info
    from src.engines.llm_engine import call_llm_engine
    inst = get_institution_info()
    proj = get_project_info()
    inst_name = inst.get("name", "项目单位")
    inst_dept = inst.get("department", "")
    proj_name = proj.get("name", "项目设计")

    prompt = f"""请为{proj_name}撰写致谢（约300-500字）。

要求：
1. 感谢指导教师（{student.advisor}）的悉心指导
2. 感谢团队成员、合作方
3. 提及本设计使用的大模型与多模态AI技术辅助
4. 感情真挚，避免过度夸张
5. 不使用夸张的形容词堆砌

只输出致谢正文。"""
    result = call_llm_engine(
        prompt=prompt,
        system_prompt="你是一个致谢文本撰写工具。撰写简洁真挚的致谢，不要编造具体的人物对话、故事或场景。感情真挚但避免夸张和虚构。",
        model="deepseek-v4-pro",
    )
    fallback_inst = f"{inst_name}{inst_dept}" if inst_dept else inst_name
    return result.strip() if result else f"感谢指导教师{student.advisor}的悉心指导，感谢{fallback_inst}的支持。"
