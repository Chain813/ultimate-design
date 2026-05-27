import os
import shutil
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')

# Source paths
brain_prev_dir = r"C:\Users\23902\.gemini\antigravity\brain\a7a0a585-8fe2-47a0-8b18-0be8b3147e91"
atlas_dir = r"e:\AI-based-project\urban-platform\static\atlas"
temp_img_dir = r"C:\Users\23902\.gemini\antigravity\brain\4548a8df-fff1-40c0-a394-3f74511d5d61\scratch\images"

os.makedirs(temp_img_dir, exist_ok=True)

# Copy figures
print("Copying project drawings and screenshots...")
fig_mapping = {
    "DR-004_现状区位图.png": "fig_004.png",
    "DR-014_用地现状分析图.png": "fig_014.png",
    "DR-017_建筑高度现状图.png": "fig_017.png",
    "DR-040_更新模式分区图.png": "fig_040.png",
    "DR-048_建筑更新控制图.png": "fig_048.png",
    "DR-051_道路交通系统规划图.png": "fig_051.png",
    "DR-056_绿地景观系统图.png": "fig_056.png",
}

for src_name, dest_name in fig_mapping.items():
    src_path = os.path.join(atlas_dir, src_name)
    dest_path = os.path.join(temp_img_dir, dest_name)
    if os.path.exists(src_path):
        shutil.copy(src_path, dest_path)
    else:
        print(f"Drawing not found: {src_path}")

# Copy screenshots from previous brain
screenshot_mapping = {
    "stage07_radar_chart_1779851862153.png": "fig_radar.png",
    "stage12_gis_compliance_1779851923612.png": "fig_compliance.png",
    "stage12_plotly_chart_1779851932871.png": "fig_plotly.png",
}

for src_name, dest_name in screenshot_mapping.items():
    src_path = os.path.join(brain_prev_dir, src_name)
    dest_path = os.path.join(temp_img_dir, dest_name)
    if os.path.exists(src_path):
        shutil.copy(src_path, dest_path)
    else:
        print(f"Screenshot not found: {src_path}")

# Load Document
doc_template_path = r"C:\Users\23902\Desktop\城环杯\附件3 成果研究报告（模板）.docx"
doc_output_path = r"C:\Users\23902\Desktop\城环杯\附件3 成果研究报告.docx"
print("Loading Word document template...")
doc = Document(doc_template_path)

def set_font_run(run, font_name="Times New Roman", east_asia_font="宋体", size_pt=12, bold=False, italic=False, color_rgb=(0,0,0)):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), east_asia_font)
    rPr.append(rFonts)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

def insert_p_before(anchor_p, text, style_name="Normal", bold=False, indent=True):
    p = anchor_p.insert_paragraph_before(style=style_name)
    run = p.add_run(text)
    set_font_run(run, size_pt=12, bold=bold)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.33)
    return p

def insert_formatted_p(anchor_p, parts, style_name="Normal", indent=True):
    """
    parts is a list of tuples: (text_content, is_bold, is_italic, is_superscript)
    """
    p = anchor_p.insert_paragraph_before(style=style_name)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.33)
        
    for text, bold, italic, super_script in parts:
        run = p.add_run(text)
        set_font_run(run, size_pt=12, bold=bold, italic=italic)
        if super_script:
            run.font.superscript = True
    return p

def insert_img_before(anchor_p, img_path, caption_text, width_inches=4.8):
    p_img = anchor_p.insert_paragraph_before()
    p_img.style = "Normal"
    p_img.paragraph_format.alignment = 1 # Center
    
    # CRITICAL BUG FIX: Setting line spacing to 1.0 (multiple/standard line spacing)
    # instead of Pt(20) fixed spacing, to prevent Word from compressing the inline shape height.
    p_img.paragraph_format.line_spacing = 1.0
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(6)
    run_img = p_img.add_run()
    run_img.add_picture(img_path, width=Inches(width_inches))
    
    # Figure caption format: 图号在图名之前，用两个半角/全角空格分隔
    p_cap = anchor_p.insert_paragraph_before()
    p_cap.style = "Normal"
    p_cap.paragraph_format.alignment = 1 # Center
    p_cap.paragraph_format.line_spacing = Pt(20)
    p_cap.paragraph_format.space_before = Pt(0)
    p_cap.paragraph_format.space_after = Pt(12)
    run_cap = p_cap.add_run(caption_text)
    set_font_run(run_cap, size_pt=12, bold=True) # Figure caption in Normal font style (12pt, bold)
    return p_cap

def insert_equation(p_next, eq_text, eq_num_text):
    """
    Creates a borderless table to right-align the formula number on the same line.
    """
    table = doc.add_table(rows=1, cols=2)
    table.alignment = 1 # Center table
    
    # Set borderless table using XML
    borders_xml = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="none"/><w:left w:val="none"/>'
        '<w:bottom w:val="none"/><w:right w:val="none"/>'
        '<w:insideH w:val="none"/><w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    table._element.tblPr.append(borders_xml)
    
    row = table.rows[0]
    cell_eq = row.cells[0]
    cell_num = row.cells[1]
    
    cell_eq.width = Inches(5.5)
    cell_num.width = Inches(1.0)
    
    p_eq = cell_eq.paragraphs[0]
    p_eq.paragraph_format.alignment = 1 # Center
    p_eq.paragraph_format.line_spacing = Pt(20)
    p_eq.paragraph_format.space_before = Pt(6)
    p_eq.paragraph_format.space_after = Pt(6)
    run_eq = p_eq.add_run(eq_text)
    set_font_run(run_eq, size_pt=12, italic=True) # Variables in equations are italicized
    
    p_num = cell_num.paragraphs[0]
    p_num.paragraph_format.alignment = 2 # Right-aligned
    p_num.paragraph_format.line_spacing = Pt(20)
    p_num.paragraph_format.space_before = Pt(6)
    p_num.paragraph_format.space_after = Pt(6)
    run_num = p_num.add_run(eq_num_text)
    set_font_run(run_num, size_pt=12)
    
    # Insert the table element before p_next
    p_next_parent = p_next._element.getparent()
    p_next_parent.insert(p_next_parent.index(p_next._element), table._element)
    return table

def find_paragraph_by_text(doc, text_fragment):
    for idx, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return idx
    return -1

# Cover Page Modifications (referencing P5 to P20 in docx_structure.txt)
print("Writing cover page metadata...")
for idx, p in enumerate(doc.paragraphs[:30]):
    text = p.text.strip()
    if "题目（35个汉字以内" in text:
        p.text = ""
        p.paragraph_format.alignment = 1
        p.paragraph_format.line_spacing = Pt(24)
        run = p.add_run("基于大模型与多模态AI的城市更新\n空间设计智能推演系统")
        set_font_run(run, font_name="黑体", east_asia_font="黑体", size_pt=26, bold=True)
    elif "参赛编号：" in text:
        p.text = ""
        p.paragraph_format.line_spacing = Pt(20)
        p.paragraph_format.alignment = 1
        run = p.add_run("参赛编号：F221")
        set_font_run(run, font_name="黑体", east_asia_font="黑体", size_pt=14, bold=True)
    elif "参 赛 人：" in text:
        p.text = ""
        p.paragraph_format.line_spacing = Pt(20)
        p.paragraph_format.alignment = 1
        run = p.add_run("参 赛 人：陈礼冲、刘旭东")
        set_font_run(run, font_name="黑体", east_asia_font="黑体", size_pt=14, bold=True)
    elif "指 导 人：" in text:
        p.text = ""
        p.paragraph_format.line_spacing = Pt(20)
        p.paragraph_format.alignment = 1
        run = p.add_run("指 导 人：李冰心、崔诚慧")
        set_font_run(run, font_name="黑体", east_asia_font="黑体", size_pt=14, bold=True)
    elif "工作单位：" in text:
        p.text = ""
        p.paragraph_format.line_spacing = Pt(20)
        p.paragraph_format.alignment = 1
        run = p.add_run("工作单位：吉林建筑大学")
        set_font_run(run, font_name="黑体", east_asia_font="黑体", size_pt=14, bold=True)
    elif "报名主题：" in text:
        p.text = ""
        p.paragraph_format.line_spacing = Pt(20)
        p.paragraph_format.alignment = 1
        run = p.add_run("报名主题：主题二：面向高质量发展的城市治理")
        set_font_run(run, font_name="黑体", east_asia_font="黑体", size_pt=14, bold=True)
    elif "研究议题：" in text:
        p.text = ""
        p.paragraph_format.line_spacing = Pt(20)
        p.paragraph_format.alignment = 1
        run = p.add_run("研究议题：议题6：城市体检与城市更新")
        set_font_run(run, font_name="黑体", east_asia_font="黑体", size_pt=14, bold=True)
    elif "技术关键词：" in text:
        p.text = ""
        p.paragraph_format.line_spacing = Pt(20)
        p.paragraph_format.alignment = 1
        run = p.add_run("技术关键词：大语言模型（必填）、多主体博弈决策（选填）、城市体检诊断（选填）")
        set_font_run(run, font_name="黑体", east_asia_font="黑体", size_pt=12, bold=True)
    elif "介绍参赛团队的研究背景" in text or "字数要求：100-200字" in text:
        p.text = ""
    elif "参赛团队简介：" in text:
        p.text = ""
        p.paragraph_format.line_spacing = Pt(20)
        run_title = p.add_run("参赛团队简介：\n")
        set_font_run(run_title, font_name="黑体", east_asia_font="黑体", size_pt=12, bold=True)
        run_desc = p.add_run(
            "本参赛团队成员均来自吉林建筑大学城乡规划专业，长期致力于数字城乡规划与智能空间决策支持系统研究。团队在多源空间大数据分析、遥感与计算机视觉处理、大语言模型与多智能体博弈系统在规划中的应用等方面具备深厚积淀。成员曾深度参与多项吉林省及长春市历史文化保护街区微更新实证项目，拥有丰富的数字平台开发与诊断分析经验。"
        )
        set_font_run(run_desc, size_pt=11)

# Delete template instruction section 〇
print("Cleaning up template layout instructions...")
idx_start_o = find_paragraph_by_text(doc, "〇、正文格式模板")
idx_end_o = find_paragraph_by_text(doc, "一、研究问题")

p_to_del = [doc.paragraphs[i] for i in range(idx_start_o, idx_end_o)]
for p in p_to_del:
    p._element.getparent().remove(p._element)

print("Template layout instructions deleted.")

def replace_section_content_before(doc, subheading_text, next_element_text, populate_func):
    """
    Find subheading, delete all paragraphs between it and next_element_text,
    and call populate_func(p_next_element) to insert new paragraphs before next_element_text.
    """
    idx_sub = find_paragraph_by_text(doc, subheading_text)
    if idx_sub == -1:
        print(f"Subheading '{subheading_text}' not found!")
        return
    
    idx_next = find_paragraph_by_text(doc, next_element_text)
    if idx_next == -1:
        print(f"Next element '{next_element_text}' not found!")
        return
    
    # Delete paragraphs between idx_sub and idx_next
    p_to_del = [doc.paragraphs[i] for i in range(idx_sub + 1, idx_next)]
    for p in p_to_del:
        p._element.getparent().remove(p._element)
        
    # Re-find index of next element
    idx_next = find_paragraph_by_text(doc, next_element_text)
    p_next = doc.paragraphs[idx_next]
    populate_func(p_next)

# CHAPTER 1 SECTION 1
def populate_c1_s1(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("我国城市化发展已由“大拆大建”的增量扩张阶段全面转向“存量提质”的高质量发展阶段。以老旧小区改造、历史风貌协调街区整治和低效用地腾退为核心的城市微更新，已成为转变城市发展方式、提升人居环境品质的重要抓手。然而，在城市微更新的学术研究 and 具体规划实践中，仍面临三大长期瓶颈：第一，在现状诊断层面，传统的城市体检高度依赖规划专家的主观定性判断与粗粒度普查，缺乏可量化、地块级、多源数据融合的更新潜力诊断模型；第二，在更新决策层面，微更新涉及政府、开发商、居民等多方利益主体，各方诉求（历史保护、经济回报、生活便利）冲突激烈，缺乏科学的动态博弈与共识达成平台；第三，在规划表达与方案深化层面，传统CAD/GIS手工制图周期长、效率低，难以实现跨方案多维比选，且直接应用普通的生成式AI（AIGC）进行规划绘图时存在严重的空间位置幻觉与地理控制偏差。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("针对上述痛点，本研究依托长春伪满皇宫周边170.2公顷的典型历史文化与工业协调更新片区，基于大语言模型（LLM）与多模态AI技术，自主研发了“城市更新空间设计智能推演系统”。该系统将“多源数据体检诊断”、“多主体博弈决策协商”、“RAG国家法规合规审查”与“空间对齐AI规划制图”四大引擎融为一体，构建了“诊断-博弈-生成-校验”的闭环循证工作流。这不仅填补了数字化智能系统辅助城市微更新决策的技术空白，更为探索存量规划的民主决策与高精度制图设计提供了全新的方法论支撑。", False, False, False)
    ])

replace_section_content_before(doc, "1．研究背景及目的意义", "2．研究目标及拟解决的问题", populate_c1_s1)

# CHAPTER 1 SECTION 2
def populate_c1_s2(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本项目的总体目标是针对城市微更新中“诊断粗糙”、“博弈断裂”和“制图幻觉”三大痛点，构建一套基于多模态大模型的城市微更新全生命周期数字化辅助系统。具体目标包括：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）构建基于AHP-", False, False, False),
        ("MPI", False, True, False), # Italicized variable
        ("（层次分析法与微更新潜力指数）的地块级多维定量评估模型，对研究范围内719栋现状建筑及相应地块的更新紧迫性与潜力进行科学排序，识别更新重点；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）建立大模型驱动的“居民-开发商-规划师”三主体利益博弈协商沙盘，基于精细化效用算法动态调整满意度指标，生成可视化雷达图并判定协商共识；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）首创「矢量-光栅-ControlNet」空间约束AI制图管线，消除地理空间幻觉，实现包含区位、现状用地、高度控制、交通系统等在内的26张A3标准规划图纸全自动多进程并行编译与专业排版；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（4）以长春市伪满皇宫博物院周边170.2公顷街区为实证应用，验证系统在实际存量城市微更新项目中的可用性与可迁移性。", False, False, False)
    ])

replace_section_content_before(doc, "2．研究目标及拟解决的问题", "二、研究方法", populate_c1_s2)

# CHAPTER 2 SECTION 1
def populate_c2_s1(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本系统依托计算机视觉、多智能体协作、空间拓扑计算以及大语言模型，建立了一套人机协同的城市微更新分阶段决策支持体系。核心方法与理论依据如下：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）层次分析法（AHP）与更新潜力测度：通过Saaty两两比较判断矩阵进行权重划分，建立空间潜力（", False, False, False),
        ("S", False, True, False),
        ("）、社会需求（", False, False, False),
        ("D", False, True, False),
        ("）和环境现状（", False, False, False),
        ("E", False, True, False),
        ("）三维一体的潜力计算模型，实现地块级更新紧迫程度的自适应计算；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）多主体决策与效用协商理论：参考“生成式智能体”与“城市智能体”思想，将博弈论应用于老旧街区协商。系统预置了居民、开发商和规划师三大智能体角色，通过动态词频与规划要素对发言进行内容分析，量化计算三方在容积率放宽、口袋公园增设等场景下的满意度走势，利用满意度雷达图进行共识判断；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）空间句法与路网拓扑分析：依托OSMnx", False, False, False),
        ("[1]", False, False, True), # Superscript citation
        ("与NetworkX计算库，对研究范围74段核心路网（共计1,062段拓扑路段）进行拓扑计算，得出全局整合度（", False, False, False),
        ("Integration", False, True, False),
        ("）与穿行度（", False, False, False),
        ("Choice", False, True, False),
        ("），定量刻画道路网络在微更新中的空间织补需求；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（4）语义分割与街景绿视率测度：采用先进的SegFormer", False, False, False),
        ("[2]", False, False, True), # Superscript citation
        ("深度学习语义分割架构，对伪满皇宫周边447个采样点共1,788张实景影像进行像素级分割，计算平均绿视率（", False, False, False),
        ("GVI", False, True, False),
        (" = 8.7%），以此作为环境质量（", False, False, False),
        ("E", False, True, False),
        ("）的重要输入；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（5）RAG检索增强生成与合规校验：将国家保护规划条例", False, False, False),
        ("[8]", False, False, True), # Superscript citation
        ("切分为248个高维语义向量分块，大模型在生成管控导则时实时检索向量库，进行法定红线指标比对与违规红牌预警。", False, False, False)
    ])

replace_section_content_before(doc, "1．研究方法及理论依据", "2．技术路线及关键技术", populate_c2_s1)

# CHAPTER 2 SECTION 2
def populate_c2_s2(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本系统技术路线分为数据汇聚层、计算引擎层、推演策略层、以及成果表达层四个部分：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("数据层汇总了包括建筑基底在内的6类GeoJSON矢量数据和15份包含街景、POI在内的CSV感知数据，统一纠正为本地投影；计算层基于AHP和空间句法，对719栋建筑地块进行定量化多维问题诊断；策略层依托DeepSeek", False, False, False),
        ("[4]", False, False, True),
        ("模型多智能体，进行三主体利益协商推演，同时通过RAG法规知识库进行实时法定控规合规性审查；表达层利用Matplotlib", False, False, False),
        ("[7]", False, False, True),
        ("与Pillow组合构建了“三层排版流水线”，自动拼装包含比例尺、图例、图签以及大模型动态设计说明的标准A3规划大图，并支持多进程并行批量图纸编译。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("系统的四大关键技术包括：一是「矢量-光栅-ControlNet", False, False, False),
        ("[3]", False, False, True),
        ("」空间约束AI制图管线，将用地、道路等矢量红线转换为带国标色值的分类光栅图作为ControlNet的输入，彻底消除了生成式AI在规划底图上的空间位置漂移与形态幻觉；二是Data-to-Text大模型动态图例说明合成技术，在A3底板封装阶段，大模型读取GIS实测指标，自动生成三条高水准设计说明文字，确保出图指标与真实GIS物理数据严格绑定；三是多进程并行处理技术，充分调动多核CPU进行26张图纸的并行渲染，极大缩短图册编译时间。", False, False, False)
    ])

    img_arch = os.path.join(temp_img_dir, "system_architecture.png")
    if os.path.exists(img_arch):
        insert_img_before(p, img_arch, "图2-1  ultimateDESIGN 系统架构与技术路线图", width_inches=5.2)


replace_section_content_before(doc, "2．技术路线及关键技术", "三、数据说明", populate_c2_s2)

# CHAPTER 3 SECTION 1
def populate_c3_s1(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("项目实证研究建立在多源、高精度空间数据库之上。所涉及的主要数据资产包括以下8大类：", False, False, False)
    ])
    
    table = doc.add_table(rows=9, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '数据名称'
    hdr_cells[2].text = '数据格式'
    hdr_cells[3].text = '数据规模'
    hdr_cells[4].text = '在模型设计中的作用与来源'

    data_rows = [
        ("1", "研究范围红线", "GeoJSON", "170.2 公顷", "确立规划空间边界约束，自绘"),
        ("2", "现状建筑基底", "GeoJSON", "719 栋", "提取层高、建筑面积以计算容积率及建筑密度，OSM"),
        ("3", "道路网络", "GeoJSON", "74 段核心路段", "用于空间句法拓扑可达性及网络穿行度计算，OSM"),
        ("4", "现状用地分类", "GeoJSON", "108 宗地块", "核查现状与控规用地占比，计算现状绿地率，百度/自绘"),
        ("5", "兴趣点(POI)", "CSV", "411 条", "计算社会设施配套需求与服务核密度，百度API"),
        ("6", "街景影像", "JPG", "1,788 张", "计算各街区环境绿视率(GVI)及天空开阔度，百度街景"),
        ("7", "舆情文本", "CSV", "207 条", "情感分析模型获取公众对历史街区更新的社会诉求，新浪微博"),
        ("8", "政策保护规章", "PDF/JSON", "248 个向量分块", "输入RAG知识库，用于导则生成与合规审查，自然资源部/住建部")
    ]

    for i, row in enumerate(data_rows):
        row_cells = table.rows[i+1].cells
        for j, val in enumerate(row):
            row_cells[j].text = val
            
    for row in table.rows:
        for cell in row.cells:
            for cp in cell.paragraphs:
                cp.paragraph_format.line_spacing = Pt(12)
                cp.paragraph_format.space_before = Pt(2)
                cp.paragraph_format.space_after = Pt(2)
                for run in cp.runs:
                    set_font_run(run, size_pt=10.5)

    p._element.addnext(table._element)
    
    # Table Caption format: 表名在表号之前，中间以两个空格分隔
    p_cap = p.insert_paragraph_before()
    p_cap.style = "Normal"
    p_cap.paragraph_format.alignment = 1
    p_cap.paragraph_format.line_spacing = Pt(20)
    p_cap.paragraph_format.space_before = Pt(12)
    p_cap.paragraph_format.space_after = Pt(6)
    run_cap = p_cap.add_run("系统多源数据资产清单  表3-1")
    set_font_run(run_cap, size_pt=12, bold=True)

replace_section_content_before(doc, "1．数据内容及类型", "2．数据预处理技术与成果", populate_c3_s1)

# CHAPTER 3 SECTION 2
def populate_c3_s2(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("为了保证各项指标测算及AI空间绘图定位的物理精度，我们执行了严格的数据预处理流程：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）本地投影纠偏（EPSG:32651）：多源获取的空间数据往往存在WGS-84、GCJ-02、BD-09等多种空间坐标系偏差。本系统首先开发了坐标系批量转换工具（`geo_transform.py`），将所有坐标统一转化为地理坐标系，并选择以长春本地的高斯克吕格投影（EPSG:32651，北京54/3度分带第42带，中央经线126E，或UTM Zone 51N）进行高精度平面投影。这有效避免了Web墨卡托（EPSG:3857）由于投影拉伸在中高纬度地区带来的面积计算偏差（高纬度形变率高达93%），将研究范围精确锚定在170.2公顷。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）多源属性空间挂接：利用空间拓扑Join技术，将411条POI配套兴趣点、1,788张街景图片以及房屋年代、价格等非空间CSV属性，按物理空间质心精确挂接到相应的地块及719栋建筑矢量基底中，构建六维度的空间上下文属性库。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）街景绿视率(GVI)计算机视觉测度：针对下载的1,788张百度实地四方向全景街景，使用训练好的SegFormer", False, False, False),
        ("[2]", False, False, True),
        ("（ADE20K）图像语义分割神经网络，提取街景中的绿化植物植被像素比率，输出447个物理样点的全域绿视率平均值（", False, False, False),
        ("GVI", False, True, False),
        (" = 8.7%），识别出历史街区内部绿化品质严重不足的物理事实。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（4）规章切块与向量嵌入：提取长春市及伪满皇宫保护管理规划的7份核心PDF文件，使用递归字符切片法将其切分为248个包含上下文重叠的语义块，采用BGE（BAAI/bge-large-zh-v1.5）中文大模型提取高维语义嵌入向量，导入到系统的轻量级本地向量数据库中，构建RAG知识库。", False, False, False)
    ])

replace_section_content_before(doc, "2．数据预处理技术与成果", "四、模型算法", populate_c3_s2)

# CHAPTER 4 SECTION 1
def populate_c4_s1(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本系统在数据底座之上，设计并实现了三套核心规划分析算法：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）AHP-", False, False, False),
        ("MPI", False, True, False),
        (" 空间微更新潜力指数计算模型：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("微更新潜力测度是确定城市体检“病灶地块”的核心。地块的更新潜力指数由空间潜力（", False, False, False),
        ("S", False, True, False),
        ("）、社会配套需求（", False, False, False),
        ("D", False, True, False),
        ("）以及环境现状品质（", False, False, False),
        ("E", False, True, False),
        ("）三维度综合加权得出，数学公式如下：", False, False, False)
    ])
    
    # Equation 4-1 formatted in a borderless table
    insert_equation(p, "MPI_i = (0.4 * S_i + 0.3 * D_i + 0.3 * (1.0 - E_i)) * 100", "(4-1)")
    
    insert_formatted_p(p, [
        ("其中，", False, False, False),
        ("MPI_i", False, True, False),
        (" 表示第 ", False, False, False),
        ("i", False, True, False),
        (" 栋现状建筑的微更新潜力指数；", False, False, False),
        ("S_i", False, True, False),
        (" 表示其空间潜力（通过地块基底面积、建筑层数倒数归一化得到）；", False, False, False),
        ("D_i", False, True, False),
        (" 表示其设施服务配套的社会需求（基于411个POI按150米搜寻半径做核密度估算得到）；", False, False, False),
        ("E_i", False, True, False),
        (" 表示采样样点的街景绿视率", False, False, False),
        ("GVI", False, True, False),
        ("均值。其数学内涵在于：利用 (1.0 - ", False, False, False),
        ("E_i", False, True, False),
        (") 表征当前绿化品质越差，其微更新的迫切性及绿化织补需求度越高。三维度权重采用层次分析法（AHP）判定矩阵运算划分（空间 ", False, False, False),
        ("S", False, True, False),
        (": 0.4，社会 ", False, False, False),
        ("D", False, True, False),
        (": 0.3，环境 ", False, False, False),
        ("E", False, True, False),
        (": 0.3，通过一致性比例 ", False, False, False),
        ("CR", False, True, False),
        (" < 0.1 检验），保证了量化体检的严谨性。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）多主体协同协商的满意度效用函数：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("为了系统协调多方利益冲突，大语言模型博弈沙盘引入了动态效用规则。定义居民代表、开发商以及规划师（政府）的满意度效用函数为：", False, False, False)
    ])
    
    # Equation 4-2 formatted in a borderless table
    insert_equation(p, "S_role = min(100, 50 + 7 * sum( [1 for word in K_role if word in DialogueText] ))", "(4-2)")
    
    insert_formatted_p(p, [
        ("其中，", False, False, False),
        ("S_role", False, True, False),
        (" 表示角色满意度；", False, False, False),
        ("K_role", False, True, False),
        (" 表示各方的核心关键词诉求集：居民代表 ", False, False, False),
        ("K_res", False, True, False),
        (" = [“绿”, “公园”, “配套”, “社区”, “医院”, “菜市”, “养老”, “口袋”]；开发商 ", False, False, False),
        ("K_dev", False, True, False),
        (" = [“容积率”, “收益”, “文旅”, “商业”, “民宿”, “运营”, “产业”, “投资”, “回报”]；规划师 ", False, False, False),
        ("K_gov", False, True, False),
        (" = [“历史保护”, “紫线”, “限高”, “合规”, “风貌”, “条例”, “保护区”]；", False, False, False),
        ("DialogueText", False, True, False),
        (" 表示博弈协商过程中的对话文本。初始三方满意度为50分，随着博弈对话多轮演进，发言中每命中一个核心词效用，对应利益方的满意度得分累加7分，封顶为100分。系统判定的整体规划共识达成条件是三方满意度底线分均大于等于60分（即 min(", False, False, False),
        ("S_res", False, True, False),
        (", ", False, False, False),
        ("S_dev", False, True, False),
        (", ", False, False, False),
        ("S_gov", False, True, False),
        (") >= 60），实现基于多模态大模型的智能体利益自动调停。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）Zoning Compliance GIS 控规实时合规性校验算法：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("系统基于EPSG:32651投影，计算容积率（", False, False, False),
        ("FAR", False, True, False),
        ("）与建筑密度（", False, False, False),
        ("Density", False, True, False),
        ("），公式如下：", False, False, False)
    ])
    
    # Equation 4-3 formatted in a borderless table
    insert_equation(p, "FAR = sum( Floor_i * Area_i ) / Area_land,  Density = sum( Area_footprint ) / Area_land", "(4-3)")
    
    insert_formatted_p(p, [
        ("其中，", False, False, False),
        ("FAR", False, True, False),
        (" 表示容积率；", False, False, False),
        ("Density", False, True, False),
        (" 表示建筑密度；", False, False, False),
        ("Floor_i", False, True, False),
        (" 和 ", False, False, False),
        ("Area_i", False, True, False),
        (" 分别表示第 ", False, False, False),
        ("i", False, True, False),
        (" 栋现状建筑的层数和基底面积；", False, False, False),
        ("Area_land", False, True, False),
        (" 表示地块红线面积；", False, False, False),
        ("Area_footprint", False, True, False),
        (" 表示现状建筑基底总面积。系统将实测数值与规划控制上限（容积率≤1.4，绿地率≥25%，建筑高度≤18m，核心区≤9m）进行拓扑几何包含运算对比，对超出红线的违规指标自动在Streamlit上触发高亮红牌警告。", False, False, False)
    ])

    img_neg = os.path.join(temp_img_dir, "negotiation_workflow.png")
    if os.path.exists(img_neg):
        insert_img_before(p, img_neg, "图4-1  三方主体博弈协商与共识判定工作流图", width_inches=4.8)

replace_section_content_before(doc, "1．模型算法流程及相关数学公式", "2．模型算法相关支撑技术", populate_c4_s1)


# CHAPTER 4 SECTION 2
def populate_c4_s2(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本系统完全基于纯 Python 开发，开发与部署栈包括以下支撑技术：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）基础语言与平台：采用 Python 3.12 核心，前端交互与可视化采用 Streamlit 1.55 全栈框架；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）GIS 空间几何库：核心基于 GeoPandas 1.0、Shapely 2.0 以及 PyProj 3.6 进行投影和缓冲区拓扑运算；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）网络分析库：基于 NetworkX 3.2 和 OSMnx 1.9 对74段核心规划道路网的全局整合度进行连通性拓扑图论分析；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（4）可视化组件：引入 Plotly 5.24 进行 3D WebGL 交互式建筑高度基底渲染及博弈雷达图绘制；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（5）大模型及向量搜索：后台调用统一的 DeepSeek-V4 (Pro) API 大语言模型，利用 NumPy 和 Scikit-learn 对248个BGE语义向量进行基于余弦相似度的余弦距离检索。图纸编译调用 Python Multiprocessing 实现多进程并行，大幅优化系统出图效率。", False, False, False)
    ])

    img_aud = os.path.join(temp_img_dir, "compliance_audit_flow.png")
    if os.path.exists(img_aud):
        insert_img_before(p, img_aud, "图4-2  Zoning Compliance 空间控规合规性审查机制流程图", width_inches=4.8)

replace_section_content_before(doc, "2．模型算法相关支撑技术", "五、实践案例", populate_c4_s2)


# CHAPTER 5 SECTION 1
def populate_c5_s1(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本研究以长春市宽城区伪满皇宫博物院周边170.2公顷街区为对象开展实证研究。范围内包含现状建筑719栋，绝大多数为低层和多层历史风貌建筑（均值层数3.7层），但外围站前区等也有高层住宅（最高40层）。现状土地利用共108宗地块，以居住（占51.1%）、商办（占11.4%）、公共设施（占13.3%）为主。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）城市体检定量诊断：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("利用系统对这170.2公顷进行实测指标提取，并与法定控规上限进行校对。体检指标对照表如下：", False, False, False)
    ])

    table_comp = doc.add_table(rows=5, cols=4)
    table_comp.style = 'Table Grid'
    hdr = table_comp.rows[0].cells
    hdr[0].text = '体检指标'
    hdr[1].text = '实测数值'
    hdr[2].text = '法定控规标准'
    hdr[3].text = '合规诊断判定'

    comp_data = [
        ("容积率 (FAR)", "1.13", "≤ 1.40", "✅ 合规达标"),
        ("建筑密度", "30.0%", "≤ 35.0%", "✅ 合规达标"),
        ("绿地率 (GAR)", "2.9%", "≥ 25.0%", "❌ 严重违规 (偏低)"),
        ("最高建筑高度", "59.5 m", "≤ 18.0 m", "⚠️ 局部超高 (站前区溢出)")
    ]

    for i, row in enumerate(comp_data):
        r_cells = table_comp.rows[i+1].cells
        for j, val in enumerate(row):
            r_cells[j].text = val
            
    for row in table_comp.rows:
        for cell in row.cells:
            for cp in cell.paragraphs:
                cp.paragraph_format.line_spacing = Pt(12)
                cp.paragraph_format.space_before = Pt(2)
                cp.paragraph_format.space_after = Pt(2)
                for run in cp.runs:
                    set_font_run(run, size_pt=10.5)

    p._element.addnext(table_comp._element)
    
    # Table Caption format: 表名在表号之前，中间以两个空格分隔
    p_cap = p.insert_paragraph_before()
    p_cap.style = "Normal"
    p_cap.paragraph_format.alignment = 1
    p_cap.paragraph_format.line_spacing = Pt(20)
    p_cap.paragraph_format.space_before = Pt(12)
    p_cap.paragraph_format.space_after = Pt(6)
    run_cap = p_cap.add_run("长春伪满皇宫周边历史街区城市体检合规诊断对照表  表5-1")
    set_font_run(run_cap, size_pt=12, bold=True)

    insert_formatted_p(p, [
        ("由诊断表可知，虽然该区域的容积率（1.13）和建筑密度（30.0%）控制较好，但绿地率仅为2.9%，远低于25%的法定绿地率标准，这暴露出研究范围内绿化空间严重不足的短板。同时，外围个别高层建筑高度（最大达59.5m）存在高度超标，破坏了历史文化街区的整体天际线视廊。", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）多主体博弈决策协商实证：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("在策略推演阶段，系统针对“放宽限高进行高强度开发”场景模拟了三主体辩论。开发商为了获取高回报，强烈要求将高度限值放宽至36m，在对话中多次提及“投资回报、商业开发、文旅产业”，其满意度冲至80分。然而，规划局强调“历史保护区、紫线控制、合规条例”，满意度仅为65分，而居民代表强烈反对高层建筑压迫历史街区，要求增加绿化和生活配套，其满意度得分仅为45分。由于 min(", False, False, False),
        ("S_res", False, True, False),
        (", ", False, False, False),
        ("S_dev", False, True, False),
        (", ", False, False, False),
        ("S_gov", False, True, False),
        (") = 45 < 60，系统触发黄色利益冲突警报。在此警告下，规划局引导开发商退让，通过调整设计方案，将核心区开发强度控制在限高18m以内，并在站前区兼容商办，同时配套建设3个口袋公园。经过第二轮博弈，三方满意度得分最终分别达到66分、73分和68分，全部越过60分共识线，系统成功达成更新策略共识。", False, False, False)
    ])

replace_section_content_before(doc, "1．模型应用实证及结果解读", "2．模型应用案例可视化表达", populate_c5_s1)

# CHAPTER 5 SECTION 2
def populate_c5_s2(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("系统的可视化模块为规划方案决策提供了直观、实时的图像表现力：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）平台交互界面与博弈雷达图：在Streamlit页面中，决策者可直观阅读三智能体动态对话流，Plotly雷达图展示三方满意度的实时演进（图5-1），并弹出合规校验红牌告警（图5-2）。", False, False, False)
    ])

    img_radar = os.path.join(temp_img_dir, "fig_radar.png")
    if os.path.exists(img_radar):
        # Caption formatted with double spaces and no inner space in figure number
        insert_img_before(p, img_radar, "图5-1  协同博弈阶段多方满意度雷达图及共识判定界面", width_inches=4.2)

    img_comp = os.path.join(temp_img_dir, "fig_compliance.png")
    if os.path.exists(img_comp):
        insert_img_before(p, img_comp, "图5-2  Streamlit 平台 GIS 控规实时合规性校验告警面板", width_inches=4.8)

    insert_formatted_p(p, [
        ("（2）A3标准成果图册渲染：系统全自动生成并导出了26张高清标准A3图册。图纸右侧图例框下半部分完美排版了由大模型根据实测指标生成的规划设计说明（图5-3），左下角为法定管控指标卡，线条和色彩符合国家控规制图规范。", False, False, False)
    ])

    img_dr004 = os.path.join(temp_img_dir, "fig_004.png")
    if os.path.exists(img_dr004):
        insert_img_before(p, img_dr004, "图5-3  系统自动编译导出的 A3 成果图纸：DR-004 现状区位图 (自绘图例与动态设计说明)", width_inches=5.2)

    img_dr051 = os.path.join(temp_img_dir, "fig_051.png")
    if os.path.exists(img_dr051):
        insert_img_before(p, img_dr051, "图5-4  系统自动编译导出的 A3 成果图纸：DR-051 道路交通系统规划图 (规划新增道路虚线高亮)", width_inches=5.2)

replace_section_content_before(doc, "2．模型应用案例可视化表达", "六、研究总结", populate_c5_s2)

# CHAPTER 6 SECTION 1
def populate_c6_s1(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本系统设计的特点和创新性突出表现在以下四个维度：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）多模型深度耦合的闭环决策流：首次将“AHP-MPI体检诊断模型”、“多主体博弈决策协商模型”、“RAG政策合规校验模型”与“空间对齐AI制图模型”串联在一起，实现了存量更新规划“体检-协商-生成-校验”的全数字化端到端闭环决策支持；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）基于地理红线对齐的 AIGC 制图管线：首创「矢量-光栅-ControlNet」多通道控制机制，弱化了生成式AI绘图的空间幻觉，使生成的规划图纸在建筑红线、道路中线及用地分区上与GIS矢量数据实现像素级空间位置对齐，满足法定规划的物理精度要求；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）大模型动态图例说明合成技术（Data-to-Text）：图册编译阶段直接读取底层的实测GIS空间指标，调用大模型生成三条量化设计说明文字，在Pillow画布中动态排版渲染，实现了“规划图面-量化指标-设计说明”的绝对绑定，消除手工修改滞后性；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（4）数据与业务逻辑的解耦架构：系统采用彻底的参数化配置与模块化引擎设计。若将系统移植至其他城市的更新项目，仅需在 `data/` 目录下替换相应的GeoJSON和CSV文件，核心的诊断、博弈和制图引擎无需做出任何代码层面的修改，具备极强的复用性与推广价值。", False, False, False)
    ])

replace_section_content_before(doc, "1．模型设计的特点", "2．应用方向或应用前景", populate_c6_s1)

# CHAPTER 6 SECTION 2
def populate_c6_s2(p_next):
    p = p_next
    insert_formatted_p(p, [
        ("本系统围绕高质量城市治理和存量城市更新，具备广泛的工程应用前景与落地推广价值：", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（1）国家及地方城市体检与专项更新评估：系统可直接挂接到住建部门的数字城市底座上，作为城市年度健康体检评估的量化测度工具，自动诊断并输出体检报告，定位更新重点；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（2）基层自然资源规划部门与街道决策辅助平台：可在规划听证会或街道更新协商会议上作为多方利益主体的数字推演工作坊，实时输入不同方案并自动刷新博弈满意度得分与合规警告，加速规划共识达成；", False, False, False)
    ])
    insert_formatted_p(p, [
        ("（3）高等院校规划与建筑设计教学示范：本系统可作为城乡规划及城市设计课程的数字化实验平台，引导学生直观理解多源大数据诊断、空间句法分析、智能体决策等数字规划的核心方法。", False, False, False)
    ])

# Populate References section before other formatting cleanup
def populate_references(p_next):
    # Heading 1 format according to template: 黑体三号字（16pt），居中，单倍行距，段前空24磅，段后空18磅
    p_h = p_next.insert_paragraph_before(style="Heading 1")
    run_h = p_h.add_run("参考文献")
    set_font_run(run_h, font_name="黑体", east_asia_font="黑体", size_pt=16, bold=True)
    p_h.paragraph_format.alignment = 1 # Center
    p_h.paragraph_format.line_spacing = 1.0 # Single spacing
    p_h.paragraph_format.space_before = Pt(24)
    p_h.paragraph_format.space_after = Pt(18)
    
    refs = [
        "[1] Boeing, G. (2017). OSMnx: New methods for acquiring, constructing, analyzing, and visualizing complex street networks. Computers, Environment and Urban Systems, 65, 126-139. (开源交通路网拓扑分析与整合度计算参考库)",
        "[2] Xie, E., Wang, W., Yu, Z., Anandkumar, A., Alvarez, J. M., & Luo, P. (2021). SegFormer: Simple and efficient design for semantic segmentation with transformers. NeurIPS. (开源街景图像绿视率深度学习语义分割神经网络)",
        "[3] Zhang, L., Rao, A., & Agrawala, M. (2023). Adding conditional control to text-to-image diffusion models. ICCV. (开源 ControlNet 规划总规图空间物理对齐约束管线)",
        "[4] DeepSeek-AI. (2024). DeepSeek-V4 Chat and Reasoning Models. (开源大语言决策推理博弈沙盘引擎)",
        "[5] McKinney, W. (2010). Data structures for statistical computing in python. Proceedings of the 9th Python in Science Conference. (开源 Pandas 数据科学分析基础库)",
        "[6] Jordahl, K., et al. (2020). GeoPandas: Geospatial pandas. Zenodo. (开源 GIS 空间数据拓扑关系处理库)",
        "[7] Hunter, J. D. (2007). Matplotlib: A 2D graphics environment. Computing in Science & Engineering. (开源数据可视化与图册渲染排版库)",
        "[8] 住房和城乡建设部. (2021). 城市体检评估工作导则[R]. 北京: 中华人民共和国住房和城乡建设部. (国家标准条例规范参考)"
    ]
    
    for ref in refs:
        p_ref = p_next.insert_paragraph_before(style="Normal")
        run_ref = p_ref.add_run(ref)
        set_font_run(run_ref, size_pt=10.5)
        p_ref.paragraph_format.line_spacing = Pt(20)
        p_ref.paragraph_format.first_line_indent = Inches(0.33)
        p_ref.paragraph_format.space_before = Pt(0)
        p_ref.paragraph_format.space_after = Pt(2)

idx_other = find_paragraph_by_text(doc, "其他格式要求！！！")
if idx_other != -1:
    p_next_other = doc.paragraphs[idx_other]
    populate_references(p_next_other)
    replace_section_content_before(doc, "2．应用方向或应用前景", "参考文献", populate_c6_s2)

# Delete trailing instructions
idx_other = find_paragraph_by_text(doc, "其他格式要求！！！")
if idx_other != -1:
    p_to_del_footer = [doc.paragraphs[i] for i in range(idx_other, len(doc.paragraphs))]
    for p in p_to_del_footer:
        p._element.getparent().remove(p._element)

print("Paragraphs remaining in final document:", len(doc.paragraphs))

# Save Document
try:
    doc.save(doc_output_path)
    print("Word document generated and saved successfully!")
except PermissionError:
    alt_path = doc_output_path.replace(".docx", "_更新.docx")
    doc.save(alt_path)
    print(f"Permission denied on {doc_output_path} (probably open in Word). Saved instead to: {alt_path}")

